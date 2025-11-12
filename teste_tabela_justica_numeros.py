"""
Script de teste para a Tabela 12 - Justiça em Números
Gera um documento contendo apenas essa tabela para facilitar ajustes.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Importar dados da tabela
try:
    from report_data import dados_tabela_justica_numeros
except ImportError:
    print("!!! ERRO: Não foi possível importar 'dados_tabela_justica_numeros' de 'report_data.py'")
    exit()


def set_row_height_at_least(row, height_twips):
    """ Define a altura MÍNIMA da linha usando XML (twips), permitindo expansão. """
    # Usa o valor passado sem impor mínimo adicional
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'atLeast')
    
    for existing_trHeight in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing_trHeight)
        
    trPr.append(trHeight)


def set_cell_vertical_alignment(cell, align='center'):
    """ Define o alinhamento vertical de uma célula. """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for existing_vAlign in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(existing_vAlign)
    
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def remove_all_borders(cell):
    """ Remove todas as bordas da célula (usando XML). """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # Define todas as bordas como 'none'
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tcBorders.append(border)
        border.set(qn('w:val'), 'none')


def set_cell_width(cell, width_cm):
    """ Define a largura de uma célula específica em centímetros. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(Cm(width_cm).twips)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def adicionar_tabela_justica_numeros(document, dados):
    """
    Cria a Tabela 12 - Dados estatísticos do Relatório Justiça em Números.
    
    Estrutura:
    - HEADER_MERGE: Cabeçalho principal (mescla todas colunas)
    - SUB_HEADER: Linha de anos (2019-2024)
    - SUB_HEADER_SECONDARY: Linha "Ano Base: ..."
    - DATA_ROW: Linhas de dados com zebrado alternado
    """
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_HEADER_PRINCIPAL_HEX = '44546A'    # Cinza escuro (cabeçalho principal)
    COR_HEADER_ANOS_HEX = 'EEEEEE'         # Cinza médio (linha de anos)
    COR_LINHA_ZEBRADA_HEX = 'D9D9D9'       # Cinza claro (zebrado)
    COR_BRANCO_RGB = RGBColor(255, 255, 255)
    COR_PRETO_RGB = RGBColor(0, 0, 0)
    
    TAMANHO_FONTE_PADRAO = Pt(11)
    FONTE = 'Calibri'
    NUM_COLUNAS = 7
    
    # Altura mínima obrigatória: 0.48 cm = 272 twips
    ALTURA_MINIMA_TWIPS = 272
    
    # --- ESTRUTURA E LARGURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    
    # NÃO aplicar nenhum estilo - isso permite controle total das larguras
    table.style = None
    
    # Larguras: Coluna 1 (indicador) mais larga, demais colunas iguais
    largura_col_indicador = Cm(5.5)
    largura_col_ano = Cm(2.25)
    
    # Obter elemento XML da tabela
    tbl = table._tbl
    
    # Configurar propriedades da tabela
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Definir que a tabela NÃO deve auto-ajustar (critical!)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Definir largura total da tabela
    tblW = OxmlElement('w:tblW')
    largura_total_twips = int(largura_col_indicador.twips + (largura_col_ano.twips * 6))
    tblW.set(qn('w:w'), str(largura_total_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    # Recuo negativo para tabela começar antes da margem esquerda
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Cm(-1.15).twips)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Forçar larguras via tblGrid (grid da tabela)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    # Criar novo grid com larguras específicas
    tblGrid = OxmlElement('w:tblGrid')
    
    # Coluna 1 (indicador)
    gridCol = OxmlElement('w:gridCol')
    gridCol.set(qn('w:w'), str(int(largura_col_indicador.twips)))
    tblGrid.append(gridCol)
    
    # Colunas 2-7 (anos)
    for _ in range(6):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(largura_col_ano.twips)))
        tblGrid.append(gridCol)
    
    # Inserir grid após tblPr
    tbl.insert(1, tblGrid)
    
    data_row_index = 0
    
    # --- PREENCHIMENTO DA TABELA ---
    for i, row_data_full in enumerate(dados):
        tipo = row_data_full[0]
        row_data = row_data_full[1:1 + NUM_COLUNAS]
        
        row = table.add_row()
        
        # Define altura mínima obrigatória para todas as linhas
        set_row_height_at_least(row, ALTURA_MINIMA_TWIPS)
        
        # Configuração XML da linha
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        
        # SUB_HEADER e SUB_HEADER_SECONDARY devem ser mantidos juntos
        if tipo in ["HEADER_MERGE", "SUB_HEADER", "SUB_HEADER_SECONDARY"]:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0')
            trPr.append(cantSplit)
        
        cells = row.cells
        
        # --- TIPO 1: CABEÇALHO PRINCIPAL (Mescla todas as colunas) ---
        if tipo == "HEADER_MERGE":
            cell = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4]).merge(cells[5]).merge(cells[6])
            cell.text = row_data[0]
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_PRINCIPAL_HEX)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            remove_all_borders(cell)  # Remove todas as bordas
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = Pt(12)
        
        # --- TIPO 2 & 3: SUB-CABEÇALHO (Anos de edição e Ano base) ---
        elif tipo in ["SUB_HEADER", "SUB_HEADER_SECONDARY"]:
            for j in range(NUM_COLUNAS):
                cell = cells[j]
                cell.text = row_data[j]
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_HEADER_ANOS_HEX)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                set_cell_vertical_alignment(cell, 'center')
                remove_all_borders(cell)  # Remove todas as bordas
                
                # Adicionar borda inferior APENAS no SUB_HEADER_SECONDARY
                if tipo == "SUB_HEADER_SECONDARY":
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    
                    # Adicionar borda inferior
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '4')
                    bottom.set(qn('w:color'), '000000')
                    tcBorders.append(bottom)
                
                p = cell.paragraphs[0]
                # Primeira coluna: alinhamento à esquerda
                # Demais colunas: centralizado
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                
                run = p.runs[0]
                run.font.color.rgb = COR_PRETO_RGB
                # SUB_HEADER em negrito, SUB_HEADER_SECONDARY em fonte normal
                run.bold = True
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
        
        # --- TIPO 4: LINHAS DE DADOS (TODAS com zebrado alternado) ---
        elif tipo == "DATA_ROW":
            data_row_index += 1
            
            for j in range(NUM_COLUNAS):
                cell = cells[j]
                cell.text = row_data[j]
                
                set_cell_vertical_alignment(cell, 'center')
                remove_all_borders(cell)  # Remove todas as bordas
                
                p = cell.paragraphs[0]
                # Primeira coluna: alinhamento à esquerda
                # Demais colunas: centralizado
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                
                run = p.runs[0]
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = False
                
                # Zebrado em TODAS as linhas de dados (linhas ímpares = cinza)
                if data_row_index % 2 == 1:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA_HEX)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # --- LEGENDA/FONTE ---
    p_legenda = document.add_paragraph(style='Normal')
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_legenda.paragraph_format.space_before = Pt(6)
    
    run_legenda = p_legenda.add_run(
        "Tabela 12 - Dados estatísticos do Relatório Justiça em Números – Edições 2019 a 2024/CNJ.\n"
        "Legenda: s/d = Dados não encontrados no Relatório Justiça em Números do Período.\n"
        "(*) O indicador considera: número de servidores(as) (efetivos(as), requisitados(as), cedidos(as) e "
        "comissionados(as) sem vínculo efetivo); e número de trabalhadores(as) auxiliares (terceirizados(as), "
        "estagiários(as), juízes(as) leigos(as) e conciliadores(as)."
    )
    run_legenda.bold = False
    run_legenda.font.name = FONTE
    run_legenda.font.size = Pt(8)
    
    p_legenda.paragraph_format.space_after = Pt(30)


def main():
    """Função principal que cria o documento de teste."""
    print("=" * 60)
    print("TESTE: Tabela 12 - Justiça em Números")
    print("=" * 60)
    
    # Criar documento
    document = Document()
    
    # Configurar margens
    section = document.sections[0]
    section.top_margin = Cm(3.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    print(f"✓ Documento criado com margens configuradas")
    
    # Adicionar título
    titulo = document.add_heading('Teste - Tabela 12: Justiça em Números', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar a tabela
    print(f"✓ Adicionando tabela com {len(dados_tabela_justica_numeros)} linhas...")
    adicionar_tabela_justica_numeros(document, dados_tabela_justica_numeros)
    
    # Salvar documento
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    CAMINHO_SAIDA = os.path.join(SCRIPT_DIR, "export", "Teste_Tabela_Justica_Numeros.docx")
    
    # Criar pasta export se não existir
    os.makedirs(os.path.dirname(CAMINHO_SAIDA), exist_ok=True)
    
    document.save(CAMINHO_SAIDA)
    
    print(f"✓ Documento salvo em: {CAMINHO_SAIDA}")
    print("=" * 60)
    print("✅ TESTE CONCLUÍDO COM SUCESSO!")
    print("=" * 60)


if __name__ == "__main__":
    main()

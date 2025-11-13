import json
import re 
import os 
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT 
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- 1. IMPORTA√á√ÉO DOS DADOS EXTERNOS ---
try:
    from report_data import dados_tabela_atos, dados_tabela_areas, dados_tabela_estrutura, dados_tabela_comarcas, dados_tabela_nucleos, dados_tabela_processos, dados_tabela_julgamentos, dados_tabela_acervo, dados_tabela_orcamento, TITULO_TABELA_ORCAMENTO, dados_tabela_orcamento_acao, TITULO_TABELA_ORCAMENTO_ACAO, dados_tabela_orcamento_2025, dados_tabela_cidades, dados_tabela_justica_numeros, MAPA_IMAGENS
except ImportError:
    print("!!! ERRO CR√çTICO: N√£o foi poss√≠vel encontrar o arquivo 'report_data.py'.")
    print("!!! Certifique-se de que 'report_data.py' est√° no mesmo diret√≥rio.")
    exit()

# --- 2. DEFINI√á√ÉO DAS EXPRESS√ïES REGULARES ---
# Pattern para sum√°rio: captura n√∫mero COM ponto inclu√≠do (se houver)
PATTERN_SUMARIO = r'^\s*(\d+(?:\.\d+)*\.?)\s+(.+?)(?:\s*\.{2,}\s*\d+\s*)?$'
# Pattern atualizado: captura o n√∫mero COM ponto opcional no final
# Para t√≠tulos de n√≠vel 1: "1." captura "1."
# Para subt√≠tulos: "3.10" captura "3.10" (sem ponto final adicional)
PATTERN_CONTEUDO = r'^\s*(\d+(?:\.\d{1,2})*\.?)\s+([A-Z√Å√Ä√Ç√É√â√ä√ç√ì√î√ï√ö√á].*)$'
PATTERN_LEGENDA = r'^(Figura|Gr√°fico)\s+\d+' 

# --- 3. CARREGAR MAPEAMENTO DE GR√ÅFICOS ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MAPEAMENTO_GRAFICOS_PATH = os.path.join(SCRIPT_DIR, "mapeamento_graficos_completo.json")
MAPEAMENTO_GRAFICOS = {}

try:
    if os.path.exists(MAPEAMENTO_GRAFICOS_PATH):
        with open(MAPEAMENTO_GRAFICOS_PATH, 'r', encoding='utf-8') as f:
            MAPEAMENTO_GRAFICOS = json.load(f)
        print(f"‚úÖ Mapeamento de gr√°ficos carregado: {len(MAPEAMENTO_GRAFICOS)} entradas")
    else:
        print(f"‚ö†Ô∏è  AVISO: Arquivo de mapeamento n√£o encontrado: {MAPEAMENTO_GRAFICOS_PATH}")
        print("‚ö†Ô∏è  Os gr√°ficos ser√£o buscados apenas no MAPA_IMAGENS do report_data.py")
except Exception as e:
    print(f"‚ö†Ô∏è  ERRO ao carregar mapeamento de gr√°ficos: {e}")
    print("‚ö†Ô∏è  Os gr√°ficos ser√£o buscados apenas no MAPA_IMAGENS do report_data.py")

def buscar_caminho_grafico(legenda_chave):
    """
    Busca o caminho completo de um gr√°fico usando o mapeamento autom√°tico.
    
    Fluxo:
    1. Recebe legenda completa do Conteudo_Fonte.docx (ex: "Gr√°fico 11 - Percentual de...")
    2. Extrai apenas "Gr√°fico X" da legenda
    3. Consulta dicionario_graficos.json para encontrar o gr√°fico original (ex: "Gr√°fico 78")
    4. Busca o arquivo extra√≠do correspondente ao gr√°fico original
    
    Args:
        legenda_chave (str): Legenda completa ou simples (ex: "Gr√°fico 11 - T√≠tulo..." ou "Gr√°fico 11")
    
    Returns:
        str or None: Caminho absoluto do arquivo ou None se n√£o encontrado
    """
    # Extrair apenas "Gr√°fico X" da legenda completa e normalizar o n√∫mero
    # Ex: "Gr√°fico 01 - Percentual de..." ‚Üí "Gr√°fico 1"
    # Ex: "Gr√°fico 11 - Percentual de..." ‚Üí "Gr√°fico 11"
    match = re.match(r'^Gr√°fico\s+(\d+)', legenda_chave, re.IGNORECASE)
    if match:
        numero = int(match.group(1))  # Converte para int para remover zeros √† esquerda
        grafico_simples = f"Gr√°fico {numero}"
    else:
        grafico_simples = legenda_chave
    
    if grafico_simples in MAPEAMENTO_GRAFICOS:
        info = MAPEAMENTO_GRAFICOS[grafico_simples]
        
        # Verificar se o gr√°fico foi mapeado com sucesso
        if info.get("status") == "encontrado" and info.get("caminho_completo"):
            grafico_original = info.get("grafico_original", "")
            caminho = info["caminho_completo"]
            
            # Converter caminho relativo para absoluto se necess√°rio
            if not os.path.isabs(caminho):
                caminho = os.path.join(SCRIPT_DIR, caminho)
            
            if os.path.exists(caminho):
                print(f"   üìä {grafico_simples} ‚Üí {grafico_original} ‚Üí {os.path.basename(caminho)}")
                return caminho
            else:
                print(f"   ‚ö†Ô∏è  Arquivo mapeado n√£o encontrado: {caminho}")
        else:
            # Gr√°fico n√£o encontrado ou n√∫mero inv√°lido
            status = info.get("status", "desconhecido")
            grafico_original = info.get("grafico_original", "N/A")
            
            if status == "numero_invalido":
                print(f"   ‚ö†Ô∏è  {grafico_simples}: N√∫mero inv√°lido no dicion√°rio ('{grafico_original}')")
            else:
                print(f"   ‚ö†Ô∏è  {grafico_simples} ‚Üí {grafico_original}: Status '{status}'")
    else:
        print(f"   ‚ö†Ô∏è  {grafico_simples}: N√£o encontrado no mapeamento")
    
    return None

def aplicar_recuo_paragrafo(paragrafo, recuo_cm):
    """
    Aplica recuo horizontal (indenta√ß√£o) a um par√°grafo usando XML.
    
    Args:
        paragrafo: Objeto par√°grafo do python-docx
        recuo_cm (float): Recuo em cent√≠metros (pode ser negativo)
    
    Exemplo:
        aplicar_recuo_paragrafo(p, -1.15)  # Recuo de -1.15cm (para esquerda)
        aplicar_recuo_paragrafo(p, 0.5)     # Recuo de 0.5cm (para direita)
    """
    pPr = paragrafo._element.get_or_add_pPr()
    
    # Criar elemento de indenta√ß√£o
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    
    # Aplicar recuo √† esquerda (left indent)
    recuo_twips = int(Cm(recuo_cm).twips)
    ind.set(qn('w:left'), str(recuo_twips))

# --- 4. DADOS BRUTOS (HARDCODED) ---
pass

# --- 4. FUN√á√ïES AUXILIARES (PAGINA√á√ÉO, ALINHAMENTO, XML) ---
def configurar_margens(documento, superior_cm, esquerda_cm, direita_cm, inferior_cm):
    """ Define as margens da se√ß√£o principal do documento em cent√≠metros. """
    # Assume que estamos trabalhando na primeira se√ß√£o do documento
    section = documento.sections[0]
    
    # Configurar tamanho do papel A4
    section.page_width = Cm(21.0)   # Largura A4
    section.page_height = Cm(29.7)  # Altura A4
    
    section.top_margin = Cm(superior_cm)
    section.left_margin = Cm(esquerda_cm)
    section.right_margin = Cm(direita_cm)
    section.bottom_margin = Cm(inferior_cm)
    
    # Configurar dist√¢ncias do cabe√ßalho e rodap√©
    section.header_distance = Cm(1.0)   # Da borda: 1 cm
    section.footer_distance = Cm(1.25)  # Da borda: 1,25 cm
    
    print(f"Tamanho do papel: A4 (21,0 x 29,7 cm)")
    print(f"Margens definidas: Superior={superior_cm}cm, Esquerda={esquerda_cm}cm.")
    print(f"Cabe√ßalho: 1.0cm da borda, Rodap√©: 1.25cm da borda")


def set_row_height_at_least(row, height_twips):
    """ Define a altura M√çNIMA da linha usando XML (twips), permitindo expans√£o. 
    Esta vers√£o aplica a altura m√≠nima obrigat√≥ria de 0.6 cm (340 twips). """
    # IMPORTANTE: 0.6 cm = 340 twips (altura m√≠nima obrigat√≥ria)
    ALTURA_MINIMA_OBRIGATORIA = 340
    
    # Garante que a altura n√£o seja menor que 0.6 cm
    altura_final = max(height_twips, ALTURA_MINIMA_OBRIGATORIA)
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(altura_final))
    trHeight.set(qn('w:hRule'), 'atLeast') # Define a regra como "At Least"
    
    for existing_trHeight in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing_trHeight)
        
    trPr.append(trHeight)


def set_row_height_flexible(row, height_twips):
    """ Define a altura M√çNIMA da linha usando XML (twips), permitindo expans√£o.
    Esta vers√£o N√ÉO aplica altura m√≠nima obrigat√≥ria - aceita qualquer valor. """
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'atLeast') # Define a regra como "At Least"
    
    for existing_trHeight in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing_trHeight)
        
    trPr.append(trHeight)


def set_row_height_exact(row, height_twips):
    """ Define a altura exata da linha usando XML (twips). """
    # IMPORTANTE: 0.6 cm = 340 twips (altura m√≠nima obrigat√≥ria)
    ALTURA_MINIMA_OBRIGATORIA = 340
    
    # Garante que a altura n√£o seja menor que 0.6 cm
    altura_final = max(height_twips, ALTURA_MINIMA_OBRIGATORIA)
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(altura_final))
    trHeight.set(qn('w:hRule'), 'exact') 
    
    for existing_trHeight in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing_trHeight)
        
    trPr.append(trHeight)


def set_cell_bottom_border(cell):
    """ Adiciona uma borda inferior s√≥lida (preta, 0.5pt) a uma c√©lula espec√≠fica. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    bottom_border = OxmlElement('w:bottom') 
    bottom_border.set(qn('w:val'), 'single') 
    bottom_border.set(qn('w:sz'), '4')       
    bottom_border.set(qn('w:color'), '000000') 
    
    tcBorders.append(bottom_border)
    tcPr.append(tcBorders)


def set_group_top_border(cell):
    """ Adiciona uma borda superior s√≥lida (preta, 0.5pt) a uma c√©lula espec√≠fica. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    top_border = OxmlElement('w:top') 
    top_border.set(qn('w:val'), 'single') 
    top_border.set(qn('w:sz'), '4')       
    top_border.set(qn('w:color'), '000000') 
    
    tcBorders.append(top_border)
    tcPr.append(tcBorders)


def set_cell_all_borders(cell):
    """ Aplica bordas s√≥lidas em todas as dire√ß√µes da c√©lula (usando XML). """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for side in ['w:top', 'w:bottom', 'w:left', 'w:right']:
        border = OxmlElement(side) 
        border.set(qn('w:val'), 'single') 
        border.set(qn('w:sz'), '4')       
        border.set(qn('w:color'), '000000') 
        tcBorders.append(border)
    
    for existing_borders in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing_borders)
    tcPr.append(tcBorders)


def remove_all_borders(cell):
    """ Remove todas as bordas da c√©lula (usando XML). """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # Define todas as bordas como 'none'
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tcBorders.append(border)
        border.set(qn('w:val'), 'none')


def set_cell_width(cell, width_cm):
    """ Define a largura de uma c√©lula espec√≠fica em cent√≠metros. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(Cm(width_cm).twips)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def add_page_number(footer):
    """ Adiciona o c√≥digo de campo de pagina√ß√£o (PAGE) a um rodap√© (footer). """
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE' 
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run_begin = p.add_run()
    run_begin.element.append(fldChar_begin)
    run_instr = p.add_run()
    run_instr.element.append(instrText)
    run_end = p.add_run()
    run_end.element.append(fldChar_end)


def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma c√©lula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) 
    tcPr.append(vAlign)


def limpar_espacamento_lista(paragraph):
    """ Remove o espa√ßamento extra de um par√°grafo de lista (Tabela 01) usando XML. """
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '240') 
    spacing.set(qn('w:lineRule'), 'auto')
    
    if hasattr(pPr, 'spacing'):
        try:
            pPr.remove(pPr.spacing)
        except:
            pass
    pPr.append(spacing)

# --- 5. FUN√á√ïES DE PROCESSAMENTO E CRIA√á√ÉO ---

def extrair_sumario_para_json(caminho_arquivo_docx, pattern_regex):
    """L√™ o DOCX (Sum√°rio) e extrai t√≠tulos (Gerador de Chave Limpa)."""
    try:
        documento_sumario = Document(caminho_arquivo_docx)
    except Exception as e:
        print(f"!!! ERRO CR√çTICO: N√£o foi poss√≠vel carregar o arquivo: {caminho_arquivo_docx}")
        print(f"!!! Detalhe do Erro: {e}")
        return [] 

    estrutura_do_relatorio = [] 

    for paragrafo in documento_sumario.paragraphs:
        texto_completo = paragrafo.text
        texto_limpo = texto_completo.replace('\xa0', ' ').strip().replace('SUM√ÅRIO', '')
        
        if not texto_limpo:
            continue
        
        match = re.search(pattern_regex, texto_limpo, re.IGNORECASE)
        
        if match:
            prefixo_completo = match.group(1).strip()
            texto_titulo = match.group(2).strip()
            
            # Remover ponto final antes de calcular o level
            prefixo_sem_ponto_final = prefixo_completo.rstrip('.')
            level = len(prefixo_sem_ponto_final.split('.'))
            
            # Usar o prefixo normalizado (sem ponto final) na chave
            texto_final_com_numero = f"{prefixo_sem_ponto_final} {texto_titulo}"

            if level >= 1:
                estrutura_do_relatorio.append({
                    "tipo": "TITULO",
                    "level": level,
                    "texto": texto_final_com_numero 
                })
    return estrutura_do_relatorio


def extrair_conteudo_mapeado(caminho_arquivo_docx, pattern_titulo, pattern_legenda):
    """
    L√™ o DOCX de conte√∫do, mapeia par√°grafos e identifica marcadores.
    (Atualizado para Tabela 05 e listas numeradas/marcadores)
    """
    print(f"Iniciando extra√ß√£o de conte√∫do de: {caminho_arquivo_docx}")
    
    conteudo_mapeado = {}
    chave_titulo_atual = None
    modo_lista = None  # Controla se estamos dentro de uma lista: None, 'NUMERADA' ou 'MARCADORES'
    itens_lista_temporaria = []  # Armazena itens enquanto a lista est√° sendo constru√≠da
    
    try:
        documento_conteudo = Document(caminho_arquivo_docx)
    except Exception as e:
        print(f"!!! ERRO CR√çTICO: N√£o foi poss√≠vel carregar o arquivo de CONTE√öDO: {caminho_arquivo_docx}")
        print(f"!!! Detalhe do Erro: {e}")
        return {} 

    for paragrafo in documento_conteudo.paragraphs:
        texto = paragrafo.text.replace('\xa0', ' ').strip()
        
        if not texto:
            continue 

        # Se estamos dentro de uma lista de marcadores ou numerada, n√£o verificamos se √© t√≠tulo
        # Primeiro verificamos os marcadores de controle de lista
        if texto == "[INICIAR_LISTA_NUMERICA]":
            modo_lista = 'NUMERADA'
            itens_lista_temporaria = []
            continue
            
        elif texto == "[FINALIZAR_LISTA_NUMERICA]":
            if modo_lista == 'NUMERADA' and itens_lista_temporaria:
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "LISTA_NUMERADA",
                    "itens": itens_lista_temporaria.copy()
                })
            modo_lista = None
            itens_lista_temporaria = []
            continue
        
        elif texto == "[INICIAR_LISTA_MARCADORES]":
            modo_lista = 'MARCADORES'
            itens_lista_temporaria = []
            continue
            
        elif texto == "[FINALIZAR_LISTA_MARCADORES]":
            if modo_lista == 'MARCADORES' and itens_lista_temporaria:
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "LISTA_MARCADORES",
                    "itens": itens_lista_temporaria.copy()
                })
            modo_lista = None
            itens_lista_temporaria = []
            continue

        # Se estamos em modo de lista, adiciona o texto √† lista sem verificar se √© t√≠tulo
        if modo_lista in ['NUMERADA', 'MARCADORES']:
            itens_lista_temporaria.append(texto)
            continue

        # Agora sim, verifica se √© um t√≠tulo (apenas se N√ÉO estivermos em modo de lista)
        match_titulo = re.search(pattern_titulo, texto, re.IGNORECASE)
        
        if match_titulo:
            prefixo = match_titulo.group(1).strip()
            titulo = match_titulo.group(2).strip()
            
            # Valida√ß√£o de t√≠tulo v√°lido:
            # DEVE ter pelo menos um ponto (seja "1." ou "3.10")
            if '.' not in prefixo:
                # SEM ponto = par√°grafo (ex: 436, 810)
                if chave_titulo_atual:
                    conteudo_mapeado[chave_titulo_atual].append({
                        "tipo": "PARAGRAFO",
                        "texto": texto
                    })
                continue
            
            # Verifica se os segmentos t√™m no m√°ximo 2 d√≠gitos
            segmentos = prefixo.replace('.', ' ').split()  # Remove pontos e divide
            if any(len(seg) > 2 for seg in segmentos):
                # N√∫mero como 3.123 ou 1.436 - trata como par√°grafo
                if chave_titulo_atual:
                    conteudo_mapeado[chave_titulo_atual].append({
                        "tipo": "PARAGRAFO",
                        "texto": texto
                    })
                continue
            
            # Normalizar o prefixo: remover ponto final para evitar incompatibilidade
            # "1." vira "1", "3.10" permanece "3.10"
            prefixo_normalizado = prefixo.rstrip('.')
            chave_titulo_atual = f"{prefixo_normalizado} {titulo}" 
            conteudo_mapeado[chave_titulo_atual] = []
            
        elif chave_titulo_atual:
            
            if texto == "[INSERIR_TABELA_ATOS_NORMATIVOS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_ATOS",
                    "dados": dados_tabela_atos 
                })
            
            elif texto == "[INSERIR_TABELA_AREAS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_AREAS",
                    "dados": dados_tabela_areas
                })
            
            elif texto == "[INSERIR_TABELA_ESTRUTURA]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_ESTRUTURA",
                    "dados": dados_tabela_estrutura
                })

            elif texto == "[INSERIR_TABELA_COMARCAS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_COMARCAS",
                    "dados": dados_tabela_comarcas
                })
            
            elif texto == "[INSERIR_TABELA_NUCLEOS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_NUCLEOS",
                    "dados": dados_tabela_nucleos
                })
                
            elif texto == "[INSERIR_TABELA_PROCESSOS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_PROCESSOS",
                    "dados": dados_tabela_processos
                })

            elif texto == "[INSERIR_TABELA_JULGAMENTOS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_JULGAMENTOS",
                    "dados": dados_tabela_julgamentos
                })

            elif texto == "[INSERIR_TABELA_ACERVO]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_ACERVO",
                    "dados": dados_tabela_acervo
                })

            elif texto == "[INSERIR_TABELA_ORCAMENTO]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_ORCAMENTO",
                    "dados": dados_tabela_orcamento
                })

            # NOVO GATILHO ADICIONADO AQUI
            elif texto == "[INSERIR_TABELA_ORCAMENTO_ACAO]":
                conteudo_mapeado[chave_titulo_atual].append({
                "tipo": "TABELA_ORCAMENTO_ACAO", # Novo tipo
                "dados": dados_tabela_orcamento_acao
                })

            # NOVO GATILHO PARA TABELA 11 CONJUNTA
            elif texto == "[INSERIR_TABELA_ORCAMENTO_CONJUNTO]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_ORCAMENTO_CONJUNTO", # Novo tipo
                    "dados": dados_tabela_orcamento_2025
                    })

            # >>> NOVA CHAVE DE QUEBRA DE P√ÅGINA ADICIONADA AQUI <<<
            elif texto == "[INSERIR_QUEBRA_PAGINA]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "QUEBRA_PAGINA", 
                    "dados": None
                })

            elif re.search(pattern_legenda, texto, re.IGNORECASE):
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "FIGURA",
                    "legenda_completa": texto 
                })
            
            # >>> NOVO: TEXTO COM DESTAQUE (come√ßa com #) <<<
            elif texto.startswith("#"):
                texto_sem_hashtag = texto[1:].strip()  # Remove o # e espa√ßos
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TEXTO_DESTAQUE",
                    "texto": texto_sem_hashtag
                })

            # ... (Dentro da fun√ß√£o extrair_conteudo_mapeado)
            
            # NOVO GATILHO PARA TABELA 12 (CIDADES)
            elif texto == "[INSERIR_TABELA_CIDADES]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_CIDADES", 
                    "dados": dados_tabela_cidades
                })

             # NOVO GATILHO PARA TABELA 12 (CIDADES)
            elif texto == "[INSERIR_TABELA_JUSTICA_NUMEROS]":
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "TABELA_JUSTICA_NUMEROS", 
                    "dados": dados_tabela_justica_numeros
                })

  
            else:
                # Se n√£o √© nenhum marcador especial, adiciona como par√°grafo normal
                conteudo_mapeado[chave_titulo_atual].append({
                    "tipo": "PARAGRAFO",
                    "texto": texto
                })

    print(f"Extra√ß√£o de conte√∫do conclu√≠da. {len(conteudo_mapeado)} t√≠tulos mapeados.")
    print("=== DEBUG: Chaves criadas no conteudo_mapeado ===")
    for chave in sorted(conteudo_mapeado.keys()):
        print(f"  '{chave}' -> {len(conteudo_mapeado[chave])} blocos")
    print("=" * 50)
    return conteudo_mapeado


def customizar_estilos_titulo(documento):
    """Aplica formata√ß√£o personalizada (Copiado da sua base)."""
    
    style_h1 = documento.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Calibri' 
    font_h1.size = Pt(18) 
    font_h1.color.rgb = RGBColor(162, 22, 18) 
    font_h1.all_caps = True 
    font_h1.bold = True
    p_format_h1 = style_h1.paragraph_format
    p_format_h1.line_spacing = 1.15
    p_format_h1.left_indent = Cm(0)  # Sem recuo no estilo (aplicaremos individualmente)
    p_format_h1.space_before = Pt(0)
    p_format_h1.space_after = Pt(8)
    
    style_h2 = documento.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Calibri' 
    font_h2.size = Pt(16)
    font_h2.color.rgb = RGBColor(162, 22, 18) 
    font_h2.bold = True
    font_h2.all_caps = False
    p_format_h2 = style_h2.paragraph_format
    p_format_h2.line_spacing = 1.15
    p_format_h2.space_before = Pt(36)
    p_format_h2.space_after = Pt(8)
    
    style_h3 = documento.styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Calibri' 
    font_h3.size = Pt(15.5)
    font_h3.color.rgb = RGBColor(162, 22, 18) 
    font_h3.bold = True
    p_format_h3 = style_h3.paragraph_format
    p_format_h3.line_spacing = 1.15
    p_format_h3.space_before = Pt(0)
    p_format_h3.space_after = Pt(8)

    style_normal = documento.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(12) 

# --- COMPONENTE: TABELA 01 (ATOS) ---
def adicionar_tabela_atos(document, dados):
    """ Cria e estiliza a Tabela de Atos Normativos. """
    
    COR_CABECALHO_RGB = RGBColor(127, 127, 127)   
    COR_CABECALHO_HEX = '7F7F7F'                  
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    COR_CINZA_CLARO_HEX = 'EEEEEE'                 

    TAMANHO_FONTE_PADRAO = Pt(12) 
    FONTE = 'Calibri'
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips
    ALTURA_MINIMA_TWIPS = 340
    
    table = document.add_table(rows=1, cols=len(dados[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Alinhamento centralizado
    
    # Definir largura preferencial da tabela (17,5 cm = 9922 twips)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9922')  # 17,5 cm em twips
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    # Definir larguras espec√≠ficas das colunas
    col_widths = [Cm(4.76), Cm(12.74)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    for i, row_data in enumerate(dados):
        if i > 0:
            row = table.add_row()
        else:
            row = table.rows[0]
        
        # Define altura m√≠nima para todas as linhas (Tabela 01 - sem obrigatoriedade)
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
            
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 

        if i == 0:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
            
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            set_cell_vertical_alignment(cell, 'center') 
            cell.text = "" 
            
            lines = cell_data.split('\n')
            is_first_content_line = True

            for k, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue 

                is_list_item = line.startswith('√º')
                
                if is_first_content_line:
                    current_paragraph = cell.paragraphs[0]
                    is_first_content_line = False
                else:
                    current_paragraph = cell.add_paragraph()

                text_to_insert = line.replace('√º', '').strip() if is_list_item else line
                run = current_paragraph.add_run(text_to_insert)

                if is_list_item:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_paragraph.style = 'List Bullet' 
                    limpar_espacamento_lista(current_paragraph)
                    
                elif i == 0:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                    current_paragraph.paragraph_format.space_after = Pt(8) 
                    current_paragraph.paragraph_format.line_spacing = 1
                
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                
                if i == 0:
                    run.font.color.rgb = COR_BRANCO_RGB 
                    run.bold = True
                else:
                    run.font.color.rgb = COR_PRETO_RGB 
                    run.bold = False
            
            if i == 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_CABECALHO_HEX)
                cell._tc.get_or_add_tcPr().append(shading_elm)
            else:
                if i % 2 == 0:
                   shading_elm = OxmlElement('w:shd')
                   shading_elm.set(qn('w:fill'), COR_CINZA_CLARO_HEX)
                   cell._tc.get_or_add_tcPr().append(shading_elm)

    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    run_titulo = p_titulo_tabela.add_run("Tabela 01 - Atos Normativos referentes √† Estrutura do TJMG. Fonte: Portal TJMG")
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12) 

# --- COMPONENTE: TABELA 02 (√ÅREAS) ---
def adicionar_tabela_areas(document, dados):
    """ Cria a Tabela 02 com alinhamento esquerdo, largura de 3cm, bordas de grupo e espa√ßamento de linha 1.0. """
    
    COR_HEADER_MAIN = '7F7F7F'    
    COR_HEADER_GROUP = 'D9D9D9'  
    COR_LINHA_ZEBRADA = 'EEEEEE'
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) 
    FONTE = 'Calibri'
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips
    ALTURA_MINIMA_TWIPS = 220
    
    table = document.add_table(rows=0, cols=2)
    
    try:
        table.style = 'Normal Table'
    except KeyError:
        print("Aviso: Estilo 'Normal Table' n√£o encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    col_widths = [Cm(14.5), Cm(3.0)] 
    table.columns[0].width = col_widths[0]
    table.columns[1].width = col_widths[1]

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        col1_texto = row_data[1]
        col2_texto = row_data[2]
        
        row = table.add_row()
        
        # Define altura m√≠nima para todas as linhas (Tabela 02 - sem obrigatoriedade)
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        
        if tipo_linha == "HEADER_MAIN":
            cell1.merge(cell2)
            cell = cell1
            cell.text = col1_texto
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_MAIN)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            
        elif tipo_linha == "HEADER_GROUP_SIGLA":
            for j, cell in enumerate([cell1, cell2]):
                cell.text = col1_texto if j == 0 else col2_texto
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                run = p.runs[0]
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = True
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_group_top_border(cell) 

                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)

        elif tipo_linha == "HEADER_GROUP_MERGED":
            cell1.merge(cell2)
            cell = cell1
            cell.text = col1_texto
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.color.rgb = COR_PRETO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            
            set_group_top_border(cell) 

            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)

        elif tipo_linha == "DATA_MERGED":
            data_row_index += 1
            cell1.merge(cell2)
            cell = cell1
            cell.text = col1_texto
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.color.rgb = COR_PRETO_RGB
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO

            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)

            if data_row_index % 2 == 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                cell._tc.get_or_add_tcPr().append(shading_elm)

        elif tipo_linha == "DATA_SPLIT":
            data_row_index += 1
            
            cell1.text = col1_texto
            set_cell_vertical_alignment(cell1, 'center')
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run1 = p1.runs[0]
            run1.font.color.rgb = COR_PRETO_RGB
            run1.font.name = FONTE
            run1.font.size = TAMANHO_FONTE_PADRAO
            
            p1_format = p1.paragraph_format
            p1_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p1_format.space_before = Pt(0)
            p1_format.space_after = Pt(0)

            cell2.text = col2_texto
            set_cell_vertical_alignment(cell2, 'center')
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.runs[0]
            run2.font.color.rgb = COR_PRETO_RGB
            run2.font.name = FONTE
            run2.font.size = TAMANHO_FONTE_PADRAO
            
            p2_format = p2.paragraph_format
            p2_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p2_format.space_before = Pt(0)
            p2_format.space_after = Pt(0)
            
            if data_row_index % 2 == 0:
                for cell in [cell1, cell2]:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 02 - Principais √°reas da Secretaria do TJMG. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12) 

# --- COMPONENTE: TABELA 03 (ESTRUTURA - 1 COLUNA) ---
def adicionar_tabela_estrutura(document, dados):
    """
    Cria a Tabela 03 (Estrutura) com 1 coluna.
    """
    
    COR_HEADER_MAIN = '7F7F7F'    
    COR_HEADER_GROUP = 'D9D9D9'  
    COR_LINHA_ZEBRADA = 'EEEEEE'
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) 
    FONTE = 'Calibri'
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips
    ALTURA_MINIMA_TWIPS = 220
    
    table = document.add_table(rows=0, cols=1) 
    
    try:
        table.style = 'Normal Table'
    except KeyError:
        print("Aviso: Estilo 'Normal Table' n√£o encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    col_widths = [Cm(17.5)] 
    table.columns[0].width = col_widths[0]

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        col1_texto = row_data[1]
        
        row = table.add_row()
        
        # Define altura m√≠nima para todas as linhas (Tabela 03 - sem obrigatoriedade)
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cell = row.cells[0] 
        cell.text = col1_texto
        
        set_cell_vertical_alignment(cell, 'center')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
        
        p_format = p.paragraph_format
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        
        run = p.runs[0]
        run.font.name = FONTE
        run.font.size = TAMANHO_FONTE_PADRAO
        run.font.color.rgb = COR_PRETO_RGB
        run.bold = False

        if tipo_linha == "HEADER_MAIN":
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_MAIN)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            
        elif tipo_linha == "HEADER_GROUP_MERGED":
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            run.bold = False
            set_group_top_border(cell) 

        elif tipo_linha == "DATA_MERGED":
            data_row_index += 1
            # N√£o aplicamos zebrado nesta tabela

    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 03 - Estruturas para a Presta√ß√£o Jurisdicional na Segunda Inst√¢ncia. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(25)

# --- COMPONENTE: TABELA 04 (COMARCAS) ---
def adicionar_tabela_comarcas(document, dados):
    """
    Cria a Tabela 04 (Comarcas) com 4 colunas e cabe√ßalho mesclado.
    """
    
    COR_HEADER_MAIN = '7F7F7F'    
    COR_LINHA_ZEBRADA = 'D9D9D9' # Cor funcional para zebrado
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) 
    FONTE = 'Calibri'
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips
    ALTURA_MINIMA_TWIPS = 220
    
    table = document.add_table(rows=0, cols=4)
    table.space_after = Pt(20) 
    
    try:
        table.style = 'Normal Table'
    except KeyError:
        print("Aviso: Estilo 'Normal Table' n√£o encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    col_widths = [Cm(4.375), Cm(4.375), Cm(4.375), Cm(4.375)] 
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        row = table.add_row()
        
        # Define altura m√≠nima para todas as linhas (Tabela 04 - sem obrigatoriedade)
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cells = row.cells 

        if tipo_linha == "HEADER_MERGE_4":
            col1_texto = row_data[1]
            cell = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
            cell.text = col1_texto
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_MAIN)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO

        elif tipo_linha == "DATA_4_COL":
            data_row_index += 1
            
            for j in range(4):
                cell = cells[j]
                cell.text = row_data[j+1] 
                
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                
                run = p.runs[0]
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = False

                if data_row_index % 2 != 0: 
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 04 - Comarcas Instaladas. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_before = Pt(20)
    p_titulo_tabela.paragraph_format.space_after = Pt(30)

# --- NOVO COMPONENTE: TABELA 05 (N√öCLEOS) ---
def adicionar_tabela_nucleos(document, dados):
    """
    Cria a Tabela 05 (N√∫cleos) com 1 coluna. Come√ßa com o estilo de grupo e 
    as linhas de dados s√£o brancas (sem zebrado).
    """
    
    COR_HEADER_GROUP = 'D9D9D9'  
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) 
    FONTE = 'Calibri'
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips
    ALTURA_MINIMA_TWIPS = 220
    
    table = document.add_table(rows=0, cols=1) 
    
    try:
        table.style = 'Normal Table'
    except KeyError:
        print("Aviso: Estilo 'Normal Table' n√£o encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    col_widths = [Cm(17.5)] 
    table.columns[0].width = col_widths[0]

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        col1_texto = row_data[1]
        
        row = table.add_row()
        
        # Define altura m√≠nima para todas as linhas (Tabela 05 - sem obrigatoriedade)
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cell = row.cells[0] 
        cell.text = col1_texto
        
        set_cell_vertical_alignment(cell, 'center')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
        
        p_format = p.paragraph_format
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        
        run = p.runs[0]
        run.font.name = FONTE
        run.font.size = TAMANHO_FONTE_PADRAO
        run.font.color.rgb = COR_PRETO_RGB
        run.bold = False

        # --- TIPO 1: Cabe√ßalho de Grupo (Primeira Linha) ---
        if tipo_linha == "HEADER_GROUP_MERGED":
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            run.bold = True
            set_group_top_border(cell) 

        # --- TIPO 2: Dados (Linhas seguintes) ---
        elif tipo_linha == "DATA_MERGED":
            # Nenhuma cor de fundo √© aplicada (mantendo-se branca/transparente)
            pass

    # --- LEGENDA/FONTE (Tabela 05) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 05 - Rela√ß√£o dos N√∫cleos de Justi√ßa 4.0 da Primeira Inst√¢ncia. Fonte: Infoguia")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(30)

# --- NOVO COMPONENTE: TABELAS DE DADOS (PROCESSOS, JULGAMENTOS, ACERVO) ---
def adicionar_tabela_processos(document, dados):
    """
    Cria tabelas de dados de 7 colunas (Processos, Julgamentos, Acervo) com 
    destaque na coluna mais recente, altura fixa e bordas parciais.
    """
    
    # --- VARI√ÅVEIS DE COR E ESTILO ---
    COR_HEADER_PRINCIPAL = '7F7F7F'    
    COR_LINHA_ZEBRADA = 'D9D9D9'       
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(12) 
    FONTE = 'Calibri'
    NUM_COLUNAS = 7     
    
    # NOVAS CORES (Hexadecimais)
    COR_SUB_HEADER_COLUNA = '44546A' 
    COR_DADOS_COLUNA = 'D5DCE4'      
    COLUNA_DESTAQUE_INDEX = 5        # Coluna do ano mais recente (2024)
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips (substitui o antigo valor)
    ALTURA_MINIMA_TWIPS = 340
    
    # --- ESTRUTURA E LARGURA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS) 
    
    try:
        table.style = 'Normal Table' 
    except KeyError:
        print("Aviso: Estilo 'Normal Table' n√£o encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    # Larguras de Coluna (3.2cm + 6 * 2.133cm = ~16cm √°rea √∫til)
    col_widths = [Cm(3.2)] + [Cm(2.133)] * 6
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    data_row_index = 0 
    
    # Extrai o nome do cabe√ßalho principal e a legenda para uso din√¢mico
    titulo_principal = dados[0][1] # Ex: "Processos Distribu√≠dos"
    
    # Determina o n√∫mero da tabela dinamicamente (para a legenda)
    if "PROCESSOS" in titulo_principal:
        num_tabela = 6
        legenda_padrao = "Tabela 06 - N√∫mero de processos distribu√≠dos. Fonte: Centro de Informa√ß√µes para a Gest√£o Institucional ‚Äì CEINFO"
    elif "JULGAMENTOS" in titulo_principal:
        num_tabela = 7
        legenda_padrao = "Tabela 07 - Julgamentos realizados. Fonte: Centro de Informa√ß√µes para a Gest√£o Institucional ‚Äì CEINFO"
    elif "ACERVO" in titulo_principal:
        num_tabela = 8
        legenda_padrao = "Tabela 08 - Dados do acervo. Fonte: Centro de Informa√ß√µes para a Gest√£o Institucional ‚Äì CEINFO"
    else:
        num_tabela = 'XX'
        legenda_padrao = f"Tabela {num_tabela} - Dados Hist√≥ricos. Fonte: CEINFO"


    for i, row_data_full in enumerate(dados):
        
        tipo_linha = row_data_full[0]
        row_data = row_data_full[1:] 
        
        row = table.add_row()
        # Define altura m√≠nima (sem obrigatoriedade de 0.6cm)
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha in ["HEADER_MERGE", "SUB_HEADER"]:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cells = row.cells 

        # --- TIPO 1: CABE√áALHO PRINCIPAL (Merscla 7 Colunas) ---
        if tipo_linha == "HEADER_MERGE":
            
            cell = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4]).merge(cells[5]).merge(cells[6])
            cell.text = row_data[0]
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_PRINCIPAL)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = Pt(12) 
        
        # --- TIPO 2: SUB-CABE√áALHO (Inst√¢ncia, 2020...) ---
        elif tipo_linha == "SUB_HEADER":
            
            for j in range(NUM_COLUNAS):
                cell = cells[j]
                cell.text = row_data[j]
                
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = True
                run.font.name = FONTE
                run.font.size = Pt(12) 
                
                # --- Destaque: Coluna 2024 ---
                if j == COLUNA_DESTAQUE_INDEX:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_SUB_HEADER_COLUNA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    run.font.color.rgb = COR_BRANCO_RGB 
                
                set_cell_bottom_border(cell)

        # --- TIPO 3 & 4: LINHAS DE DADOS E TOTAL ---
        elif tipo_linha in ["DATA_ROW", "TOTAL_ROW"]:
            
            is_total_row = (tipo_linha == "TOTAL_ROW")
            if not is_total_row:
                data_row_index += 1 

            for j, cell_data in enumerate(row_data):
                cell = cells[j]
                cell.text = cell_data
                
                cell_align = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                p.alignment = cell_align
                
                run = p.runs[0]
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = is_total_row
                
                # --- L√≥gica de Sombreamento (Prioridade: Coluna > Total/Zebrado) ---
                current_shading_color = None

                if is_total_row or (data_row_index % 2 != 0): 
                    current_shading_color = COR_LINHA_ZEBRADA 
                
                if j == COLUNA_DESTAQUE_INDEX:
                    current_shading_color = COR_DADOS_COLUNA
                
                if current_shading_color:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), current_shading_color)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                # --- Fim da L√≥gica de Sombreamento ---

                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                

    # --- LEGENDA/FONTE (Tabela 06/07/08) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run(legenda_padrao)
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(30)

# --- NOVO COMPONENTE: TABELA 09 (OR√áAMENTO - 2 COLUNAS) ---
def adicionar_tabela_orcamento(document, titulo_acima, dados):
    """
    Cria a Tabela 09 (Or√ßamento - 2 colunas) com:
    - Altura fixa nos cabe√ßalhos (HEADER_MERGE e SUB_HEADER).
    - Altura M√çNIMA de 1 Pt (20 Twips) a mais nas linhas de dados.
    """
    # --- VARI√ÅVEIS DE COR E ESTILO ---
    COR_HEADER_ESCURO_HEX = '7F7F7F'    
    COR_LINHA_TOTAL_HEX = 'BFBFBF'      
    COR_BRANCO_RGB = RGBColor(255, 255, 255)
    COR_PRETO_RGB = RGBColor(0, 0, 0)
    COR_BRANCO_HEX = 'FFFFFF'
    
    TAMANHO_FONTE_PADRAO = Pt(12)
    FONTE = 'Calibri'
    NUM_COLUNAS_DADOS = 2  
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips para TODAS as linhas
    ALTURA_MINIMA_TWIPS = 340 

    # --- FUN√á√ÉO AUXILIAR DE BORDA (Aplicada em todas as dire√ß√µes) ---
    def set_cell_all_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        for side in ['w:top', 'w:bottom', 'w:left', 'w:right']:
            border = OxmlElement(side) 
            border.set(qn('w:val'), 'single') 
            border.set(qn('w:sz'), '4')       
            border.set(qn('w:color'), '000000') 
            tcBorders.append(border)
        
        for existing_borders in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(existing_borders)
        tcPr.append(tcBorders)

    # --- FUN√á√ÉO AUXILIAR PARA REMOVER TODAS AS BORDAS ---
    def remove_all_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        for existing_borders in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(existing_borders)


    # --- 1. PLOTAR O T√çTULO ACIMA DA TABELA ---
    p_titulo_acima = document.add_paragraph(style='Normal')
    p_titulo_acima.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_acima.paragraph_format.space_after = Pt(6)
    
    run_titulo = p_titulo_acima.add_run(titulo_acima)
    run_titulo.bold = True
    run_titulo.font.name = FONTE
    run_titulo.font.size = TAMANHO_FONTE_PADRAO
    run_titulo.font.color.rgb = COR_PRETO_RGB 
    
    
    # --- 2. ESTRUTURA E LARGURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS_DADOS)
    table.style = 'Table Grid'
    
    largura_total_cm = 16.0 
    table.columns[0].width = Cm(largura_total_cm * 0.60)
    table.columns[1].width = Cm(largura_total_cm * 0.40)

    legenda_padrao = "Tabela 09 - Despesa realizada por a√ß√£o (Unidade 1031). Fonte: Armaz√©m de Informa√ß√µes - BO SIAFI/MG"
    
    data_row_index = 0
    
    # --- 3. PREENCHIMENTO DA TABELA ---
    for i, row_data_full in enumerate(dados):
        tipo = row_data_full[0]
        dados_da_linha = row_data_full[1:1 + NUM_COLUNAS_DADOS] 
        
        row = table.add_row()
        
        # Define altura m√≠nima de 0.6 cm para TODAS as linhas, permitindo expans√£o
        set_row_height_at_least(row, ALTURA_MINIMA_TWIPS)
            
        if tipo == "DATA_ROW" or tipo == "TOTAL_ROW":
            data_row_index += 1
            
        for col_idx in range(NUM_COLUNAS_DADOS):
            cell = row.cells[col_idx]
            
            remove_all_borders(cell)
            set_cell_vertical_alignment(cell, 'center')
            
            p = cell.paragraphs[0]
            # Alinhamento do conte√∫do (Col 0: Left, Col 1: Center)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            p.text = dados_da_linha[col_idx]
            run = p.runs[0] 
            
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            run.font.color.rgb = COR_PRETO_RGB 

            # TIPO 1: SUB-HEADER
            if tipo == "SUB_HEADER":
                set_cell_all_borders(cell)
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_HEADER_ESCURO_HEX)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                run.font.color.rgb = COR_BRANCO_RGB
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centraliza o texto do SUB-HEADER
                
            # TIPO 2 & 3: DADOS (DATA_ROW e TOTAL_ROW)
            elif tipo in ["DATA_ROW", "TOTAL_ROW"]:
                
                set_cell_all_borders(cell)
                
                if tipo == "TOTAL_ROW":
                    run.bold = True
                    
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_TOTAL_HEX)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # --- 4. LEGENDA/FONTE FINAL ---
    p_legenda = document.add_paragraph(style='Normal')
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_legenda.paragraph_format.space_before = Pt(6)
    
    run_legenda = p_legenda.add_run(legenda_padrao)
    run_legenda.bold = False 
    run_legenda.font.name = FONTE
    run_legenda.font.size = Pt(8)
    
    p_legenda.paragraph_format.space_after = Pt(15)

# --- NOVO COMPONENTE: TABELA 11 (OR√áAMENTO CONJUNTO) ---
def adicionar_tabela_orcamento_conjunto(document, dados):
    """
    Cria a Tabela 11, combinando UO 1031 e UO 4031 em uma √∫nica tabela,
    com corre√ß√£o est√©tica e l√≥gica de altura para o texto.
    """
    # --- VARI√ÅVEIS DE COR E ESTILO ---
    COR_HEADER_ESCURO_HEX = '7F7F7F'    # Cor do GROUP_TITLE
    COR_LINHA_TOTAL_HEX = 'BFBFBF'      # Cor da linha TOTAL
    COR_BRANCO_RGB = RGBColor(255, 255, 255) 
    COR_PRETO_RGB = RGBColor(0, 0, 0)
    
    TAMANHO_FONTE_PADRAO = Pt(12)
    FONTE = 'Calibri'
    NUM_COLUNAS_DADOS = 2  
    
    # Altura m√≠nima obrigat√≥ria de 0.6 cm = 340 twips para TODAS as linhas
    ALTURA_MINIMA_TWIPS = 340

    # --- 1. ESTRUTURA E LARGURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS_DADOS)
    table.style = 'Table Grid'
    
    largura_total_cm = 16.0 
    table.columns[0].width = Cm(largura_total_cm * 0.60)
    table.columns[1].width = Cm(largura_total_cm * 0.40)

    legenda_padrao = "Tabela 11 - Or√ßamento 2024 por a√ß√£o or√ßament√°ria. Fonte: Lei Or√ßament√°ria Anual n¬∫ 24.678, de 30/12/2024. Fonte: Lei Or√ßament√°ria Anual n¬∫ 25.124, de 30/12/2024."
    
    # --- 2. PREENCHIMENTO DA TABELA ---
    for i, row_data_full in enumerate(dados):
        tipo = row_data_full[0]
        dados_da_linha = row_data_full[1:1 + NUM_COLUNAS_DADOS] 
        
        row = table.add_row()
        
        # Define altura m√≠nima de 0.6 cm para TODAS as linhas, permitindo expans√£o
        set_row_height_at_least(row, ALTURA_MINIMA_TWIPS)
            
        
        for col_idx in range(NUM_COLUNAS_DADOS):
            cell = row.cells[col_idx]
            
            remove_all_borders(cell)
            set_cell_vertical_alignment(cell, 'center')
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # --- TIPO 1: GROUP_TITLE (CORRIGIDO) ---
            if tipo == "GROUP_TITLE":
                if col_idx == 0:
                    
                    cell.merge(row.cells[1]) 
                    
                    p.text = dados_da_linha[col_idx] 
                    run = p.runs[0] 
                    
                    # Estilo: Branco/Negrito/Esquerda (Corpo do Or√ßamento)
                    run.font.name = FONTE
                    run.font.size = TAMANHO_FONTE_PADRAO
                    run.font.color.rgb = COR_BRANCO_RGB 
                    run.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    set_cell_all_borders(cell) 
                    
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_HEADER_ESCURO_HEX)
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                continue 
            
            # --- PROCESSAMENTO GERAL (Sub-Header e Dados) ---
            
            p.text = dados_da_linha[col_idx] 
            run = p.runs[0] 
            
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            run.font.color.rgb = COR_PRETO_RGB 
            run.bold = False

            # TIPO 2: SUB-HEADER (EST√âTICA CORRIGIDA)
            if tipo == "SUB_HEADER":
                set_cell_all_borders(cell)
                
                # Fundo BRANCO (nenhum sombreamento √© aplicado)
                
                # Texto PRETO, Negrito, Centralizado
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # TIPO 3 & 4: DADOS (DATA_ROW e TOTAL_ROW)
            elif tipo in ["DATA_ROW", "TOTAL_ROW"]:
                
                set_cell_all_borders(cell)
                
                # Alinhamento da coluna de DADOS: Esquerda (Col 0) ou Centro (Col 1)
                # Alinhamento do TOTAL_ROW: Centro (todas as colunas)
                if tipo == "TOTAL_ROW":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centralizado
                else: # DATA_ROW
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT

                run.font.color.rgb = COR_PRETO_RGB 
                run.bold = False

                if tipo == "TOTAL_ROW":
                    run.bold = True
                    
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_TOTAL_HEX)
                    cell._tc.get_or_add_tcPr().append(shading_elm)


    # --- 3. LEGENDA/FONTE FINAL ---
    p_legenda = document.add_paragraph(style='Normal')
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_legenda.paragraph_format.space_before = Pt(6)
    
    run_legenda = p_legenda.add_run(legenda_padrao)
    run_legenda.bold = False 
    run_legenda.font.name = FONTE
    run_legenda.font.size = Pt(8)
    
    p_legenda.paragraph_format.space_after = Pt(15)        


# --- NOVO COMPONENTE: TABELA 12 (CIDADES - 4 COLUNAS) ---
def adicionar_tabela_cidades(document, dados):
    """
    Cria a Tabela de Cidades (4 colunas) com linhas zebradas, sem t√≠tulo/legenda,
    e altura de 1,03 cm.
    """
    # --- VARI√ÅVEIS DE COR E ESTILO ---
    COR_LINHA_ZEBRADA_HEX = 'D9D9D9'  # Cinza Claro
    COR_PRETO_RGB = RGBColor(0, 0, 0)
    TAMANHO_FONTE_PADRAO = Pt(12)
    FONTE = 'Calibri'
    NUM_COLUNAS = 4 
    
    # Altura das c√©lulas: 1,03 cm = 584 Twips
    ALTURA_CELULA_TWIPS = 584

    # --- ESTRUTURA E LARGURA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.style = 'Table Grid' 

    # Aplicar recuo de 1,27 cm para alinhar com o marcador anterior
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Adicionar recuo √† esquerda (tblInd)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Cm(1.27).twips)))  # Recuo de 1,27 cm
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    # Define largura igual para as 4 colunas
    largura_coluna = Cm(14.0 / NUM_COLUNAS) 
    for col in table.columns:
        col.width = largura_coluna
    
    data_row_index = 0

    # --- PREENCHIMENTO DA TABELA ---
    for i, row_data_full in enumerate(dados):
        tipo = row_data_full[0]
        row_data = row_data_full[1:] 
        
        row = table.add_row()
        
        # Aplica altura de 1,03 cm a todas as linhas
        set_row_height_flexible(row, ALTURA_CELULA_TWIPS) 

        if tipo == "DATA_ROW":
            data_row_index += 1
        
        for col_idx in range(NUM_COLUNAS):
            cell = row.cells[col_idx]
            
            # Limpa bordas padr√£o e aplica bordas completas
            remove_all_borders(cell)
            set_cell_all_borders(cell)

            set_cell_vertical_alignment(cell, 'center')
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centraliza o texto
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            p.text = row_data[col_idx]
            run = p.runs[0] 
            
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            run.font.color.rgb = COR_PRETO_RGB 
            run.bold = False

            # L√≥gica de Zebrado (apenas para DATA_ROW)
            if tipo == "DATA_ROW" and (data_row_index % 2) != 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA_HEX)
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Adiciona espa√ßamento ap√≥s a tabela
    p_espaco = document.add_paragraph()
    p_espaco.paragraph_format.space_after = Pt(12)


def adicionar_tabela_justica_numeros(document, dados):
    """
    Cria a Tabela 12 - Dados estat√≠sticos do Relat√≥rio Justi√ßa em N√∫meros.
    
    Estrutura:
    - HEADER_MERGE: Cabe√ßalho principal (mescla todas colunas)
    - SUB_HEADER: Linha de anos (2019-2024)
    - SUB_HEADER_SECONDARY: Linha "Ano Base: ..."
    - DATA_ROW: Linhas de dados com zebrado alternado
    """
    # --- VARI√ÅVEIS DE COR E ESTILO ---
    COR_HEADER_PRINCIPAL_HEX = '44546A'    # Azul TJMG (cabe√ßalho principal)
    COR_HEADER_ANOS_HEX = 'EEEEEE'         # Cinza claro (linha de anos)
    COR_LINHA_ZEBRADA_HEX = 'D9D9D9'       # Cinza m√©dio (zebrado)
    COR_BRANCO_RGB = RGBColor(255, 255, 255)
    COR_PRETO_RGB = RGBColor(0, 0, 0)
    
    TAMANHO_FONTE_PADRAO = Pt(11)
    FONTE = 'Calibri'
    NUM_COLUNAS = 7
    
    # Altura m√≠nima obrigat√≥ria: 0.48 cm = 272 twips
    ALTURA_MINIMA_TWIPS = 272
    
    # --- ESTRUTURA E LARGURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    
    # N√ÉO aplicar nenhum estilo - isso permite controle total das larguras
    table.style = None
    
    # Larguras: Coluna 1 (indicador) mais larga, demais colunas iguais
    largura_col_indicador = Cm(5.5)
    largura_col_ano = Cm(2.25)
    
    # Obter elemento XML da tabela
    tbl = table._tbl
    
    # Configurar propriedades da tabela
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Definir que a tabela N√ÉO deve auto-ajustar (critical!)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Definir largura total da tabela
    tblW = OxmlElement('w:tblW')
    largura_total_twips = int(largura_col_indicador.twips + (largura_col_ano.twips * 6))
    tblW.set(qn('w:w'), str(largura_total_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    # Recuo negativo para tabela come√ßar antes da margem esquerda
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Cm(-1.15).twips)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # For√ßar larguras via tblGrid (grid da tabela)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    # Criar novo grid com larguras espec√≠ficas
    tblGrid = OxmlElement('w:tblGrid')
    
    # Coluna 1 (indicador)
    gridCol = OxmlElement('w:gridCol')
    gridCol.set(qn('w:w'), str(int(largura_col_indicador.twips)))
    tblGrid.append(gridCol)
    
    # Colunas 2-7 (anos)
    for _ in range(6):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(largura_col_ano.twips)))
        tblGrid.append(gridCol)
    
    # Inserir grid ap√≥s tblPr
    tbl.insert(1, tblGrid)
    
    data_row_index = 0
    
    # --- PREENCHIMENTO DA TABELA ---
    for i, row_data_full in enumerate(dados):
        tipo = row_data_full[0]
        row_data = row_data_full[1:1 + NUM_COLUNAS]
        
        row = table.add_row()
        
        # Define altura m√≠nima de 0.48 cm (272 twips) - usa fun√ß√£o flex√≠vel
        set_row_height_flexible(row, ALTURA_MINIMA_TWIPS)
        
        # Configura√ß√£o XML da linha
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        
        # Impedir quebra de linha entre p√°ginas (TODAS as linhas)
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        
        # SUB_HEADER e SUB_HEADER_SECONDARY devem ser mantidos juntos como cabe√ßalhos
        if tipo in ["HEADER_MERGE", "SUB_HEADER", "SUB_HEADER_SECONDARY"]:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        
        cells = row.cells
        
        # --- TIPO 1: CABE√áALHO PRINCIPAL (Mescla todas as colunas) ---
        if tipo == "HEADER_MERGE":
            cell = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4]).merge(cells[5]).merge(cells[6])
            cell.text = row_data[0]
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_PRINCIPAL_HEX)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            remove_all_borders(cell)  # Remove todas as bordas
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = Pt(12)
        
        # --- TIPO 2 & 3: SUB-CABE√áALHO (Anos de edi√ß√£o e Ano base) ---
        elif tipo in ["SUB_HEADER", "SUB_HEADER_SECONDARY"]:
            for j in range(NUM_COLUNAS):
                cell = cells[j]
                cell.text = row_data[j]
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_HEADER_ANOS_HEX)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                set_cell_vertical_alignment(cell, 'center')
                remove_all_borders(cell)  # Remove todas as bordas
                
                # Adicionar borda inferior APENAS no SUB_HEADER_SECONDARY
                if tipo == "SUB_HEADER_SECONDARY":
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    
                    # Adicionar borda inferior
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '4')
                    bottom.set(qn('w:color'), '000000')
                    tcBorders.append(bottom)
                
                p = cell.paragraphs[0]
                # Primeira coluna: alinhamento √† esquerda
                # Demais colunas: centralizado
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                
                run = p.runs[0]
                run.font.color.rgb = COR_PRETO_RGB
                # SUB_HEADER e SUB_HEADER_SECONDARY ambos em negrito
                run.bold = True
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
        
        # --- TIPO 4: LINHAS DE DADOS (TODAS com zebrado alternado) ---
        elif tipo == "DATA_ROW":
            data_row_index += 1
            
            for j in range(NUM_COLUNAS):
                cell = cells[j]
                cell.text = row_data[j]
                
                set_cell_vertical_alignment(cell, 'center')
                remove_all_borders(cell)  # Remove todas as bordas
                
                p = cell.paragraphs[0]
                # Primeira coluna: alinhamento √† esquerda
                # Demais colunas: centralizado
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                
                run = p.runs[0]
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = False
                
                # Zebrado em TODAS as linhas de dados (linhas √≠mpares = cinza)
                if data_row_index % 2 == 1:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA_HEX)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # --- LEGENDA/FONTE ---
    p_legenda = document.add_paragraph(style='Normal')
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_legenda.paragraph_format.space_before = Pt(6)
    
    run_legenda = p_legenda.add_run(
        "Tabela 12 - Dados estat√≠sticos do Relat√≥rio Justi√ßa em N√∫meros ‚Äì Edi√ß√µes 2019 a 2024/CNJ.\n"
        "Legenda: s/d = Dados n√£o encontrados no Relat√≥rio Justi√ßa em N√∫meros do Per√≠odo.\n"
        "(*) O indicador considera: n√∫mero de servidores(as) (efetivos(as), requisitados(as), cedidos(as) e "
        "comissionados(as) sem v√≠nculo efetivo); e n√∫mero de trabalhadores(as) auxiliares (terceirizados(as), "
        "estagi√°rios(as), ju√≠zes(as) leigos(as) e conciliadores(as)."
    )
    run_legenda.bold = False
    run_legenda.font.name = FONTE
    run_legenda.font.size = Pt(8)
    
    p_legenda.paragraph_format.space_after = Pt(30)


def aplicar_estilo_capa(paragrafo, texto, tamanho_pt):
    """Aplica o estilo de fonte Bahnschrift com um tamanho espec√≠fico."""
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(texto)
    run.font.name = 'Bahnschrift SemiCondensed' 
    run.font.size = Pt(tamanho_pt)
    run.bold = True                          

# --- 6. EXECU√á√ÉO E GERA√á√ÉO DO DOCUMENTO ---

if __name__ == "__main__":
    
    print("--- INICIANDO GERADOR DE RELAT√ìRIO COMPLETO ---")
    
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    CAMINHO_SUMARIO = os.path.join(SCRIPT_DIR, "Sumario_Modelo.docx")
    CAMINHO_CONTEUDO = os.path.join(SCRIPT_DIR, "Conteudo_Fonte.docx")
    CAMINHO_SAIDA = os.path.join(SCRIPT_DIR, "export/Relatorio_Final_Completo.docx")

    estrutura_final = extrair_sumario_para_json(CAMINHO_SUMARIO, PATTERN_SUMARIO)
    conteudo_mapeado = extrair_conteudo_mapeado(CAMINHO_CONTEUDO, PATTERN_CONTEUDO, PATTERN_LEGENDA)

    if not estrutura_final:
        print("Execu√ß√£o interrompida. 'estrutura_final' (Sum√°rio) est√° vazia.")
        exit()

    document = Document()

    configurar_margens(document, 3.0, 3.0, 2.0, 2.0) 

    customizar_estilos_titulo(document)
    footer = document.sections[0].footer
    add_page_number(footer)

    # --- IN√çCIO DA SE√á√ÉO: CRIA√á√ÉO DA CAPA ---
    print("Criando a Capa...")
    
    p_titulo1 = document.add_paragraph()
    aplicar_estilo_capa(p_titulo1, "RELAT√ìRIO DIAGN√ìSTICO DO TRIBUNAL", 20)
    p_titulo1.paragraph_format.space_before = Pt(48)
    p_titulo2 = document.add_paragraph()
    aplicar_estilo_capa(p_titulo2, "DE JUSTI√áA DO ESTADO DE MINAS GERAIS ‚Äì TJMG", 20)
    p_ano = document.add_paragraph()
    aplicar_estilo_capa(p_ano, "2025", 20)
    p_ano.paragraph_format.space_before = Pt(48)
    p_ano.paragraph_format.space_after = Pt(48)
    p_plan = document.add_paragraph()
    aplicar_estilo_capa(p_plan, "PLANEJAMENTO ESTRAT√âGICO 2021-2026", 20)
    p_plan.paragraph_format.space_after = Pt(280)
    p_setor = document.add_paragraph()
    aplicar_estilo_capa(p_setor, 'DEPLAG - TJMG', 14)
    p_data = document.add_paragraph()
    aplicar_estilo_capa(p_data, "JANEIRO DE 2025", 14)

    document.add_page_break()
    # --- FIM DA CAPA ---

    # --- IN√çCIO DA SE√á√ÉO: SUM√ÅRIO (Simplificado) ---
    document.add_heading('Sum√°rio', level=1)
    print("Criando sum√°rio est√°tico...")

    for elemento in estrutura_final:
        if elemento['tipo'] == 'TITULO':
            level = elemento['level']
            texto = elemento['texto']
            
            p = document.add_paragraph(style='Normal')
            run = p.add_run(texto)
            run.bold = True
            run.font.name = 'Calibri'
            
            p_format = p.paragraph_format
            p_format.line_spacing = 1.25 
            p_format.space_after = Pt(8) 
            
            if level == 2:
                p_format.left_indent = Inches(0.2)
            elif level == 3:
                p_format.left_indent = Inches(0.4)
            else:
                p_format.left_indent = Inches(0) 

    document.add_page_break()
    # --- FIM DO SUM√ÅRIO ---

    # --- SE√á√ÉO: CORPO DO DOCUMENTO (COM PROCESSADOR DE CONTE√öDO) ---
    print("Gerando corpo do relat√≥rio com t√≠tulos e conte√∫do...")
    print("=== DEBUG: Verificando correspond√™ncia de chaves ===")

    for elemento in estrutura_final:
        if elemento['tipo'] == 'TITULO':
            
            texto_chave = elemento['texto'] 
            level = elemento['level']
            
            print(f"\nBuscando conte√∫do para: '{texto_chave}'")
            print(f"  DEBUG: Level detectado = {level}, Type = {type(level)}")
            if texto_chave in conteudo_mapeado:
                print(f"  ‚úì ENCONTRADO: {len(conteudo_mapeado[texto_chave])} blocos")
            else:
                print(f"  ‚úó N√ÉO ENCONTRADO no conteudo_mapeado")
            
            if level == 1:
                texto_para_imprimir = texto_chave.replace(" ", ". ", 1)
            else:
                texto_para_imprimir = texto_chave
                
            # document.add_heading(texto_para_imprimir, level=level)
            p = document.add_heading(texto_para_imprimir, level=level)
            # Aplicar recuo de 1,25 cm para Heading 1 (t√≠tulos numerados)
            if level == 1:
                print(f"  DEBUG: Aplicando recuo ao t√≠tulo: '{texto_para_imprimir}'")
                # Usar XML para garantir que o recuo seja aplicado
                pPr = p._element.get_or_add_pPr()
                # Remover qualquer w:ind existente
                existing_ind = pPr.find(qn('w:ind'))
                if existing_ind is not None:
                    print(f"  DEBUG: Removendo w:ind existente com left={existing_ind.get(qn('w:left'))}")
                    pPr.remove(existing_ind)
                # Adicionar novo w:ind
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), '720')  # 720 twips = 1.27 cm
                pPr.append(ind)
                print(f"  ‚Üí Recuo aplicado via XML: 720 twips (1.27 cm)")

            titulo_chave = elemento['texto'] 
            
            if titulo_chave in conteudo_mapeado:
                
                for bloco in conteudo_mapeado[titulo_chave]:
                    
                    # --- PROCESSADOR DE BLOCOS ---
                    
                    if bloco['tipo'] == 'PARAGRAFO':
                        p = document.add_paragraph(bloco['texto'])
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_format = p.paragraph_format
                        p_format.line_spacing = 1.5
                        p_format.space_before = Pt(0)
                        p_format.space_after = Pt(8)
                    
                    # >>> NOVO PROCESSADOR DE TEXTO COM DESTAQUE <<<
                    elif bloco['tipo'] == 'TEXTO_DESTAQUE':
                        p = document.add_paragraph()
                        run = p.add_run(bloco['texto'])
                        run.bold = True
                        run.font.color.rgb = RGBColor(162, 22, 18)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_format = p.paragraph_format
                        p_format.line_spacing = 1.5
                        p_format.space_after = Pt(8)

                        # >>> NOVO PROCESSADOR DE QUEBRA DE P√ÅGINA AQUI <<<
                    elif bloco['tipo'] == 'QUEBRA_PAGINA':
                        print("--- Inserindo QUEBRA DE P√ÅGINA for√ßada ---")
                        document.add_page_break()
                    
                    # >>> NOVO PROCESSADOR DE LISTA NUMERADA <<<
                    elif bloco['tipo'] == 'LISTA_NUMERADA':
                        print(f"--- Inserindo LISTA NUMERADA com {len(bloco['itens'])} itens ---")
                        for item_texto in bloco['itens']:
                            p = document.add_paragraph(item_texto, style='List Number')
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_format = p.paragraph_format
                            p_format.line_spacing = 1.0
                            p_format.space_after = Pt(4)
                            # Adiciona recuo de par√°grafo
                            p_format.left_indent = Cm(1.27)  # Recuo de 1,27 cm
                        # Adiciona espa√ßo ap√≥s a lista
                        document.paragraphs[-1].paragraph_format.space_after = Pt(12)
                    
                    # >>> NOVO PROCESSADOR DE LISTA COM MARCADORES <<<
                    elif bloco['tipo'] == 'LISTA_MARCADORES':
                        print(f"--- Inserindo LISTA COM MARCADORES com {len(bloco['itens'])} itens ---")
                        for item_texto in bloco['itens']:
                            p = document.add_paragraph(item_texto, style='List Bullet')
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_format = p.paragraph_format
                            p_format.line_spacing = 1.5
                            p_format.space_after = Pt(4)
                            # Adiciona recuo de par√°grafo
                            p_format.left_indent = Cm(1.27)  # Recuo de 1,27 cm
                        # Adiciona espa√ßo ap√≥s a lista
                        document.paragraphs[-1].paragraph_format.space_after = Pt(12)
                    
                    elif bloco['tipo'] == 'FIGURA':
                        legenda_completa = bloco['legenda_completa']
                        
                        partes = legenda_completa.split("Fonte:")
                        legenda_chave = partes[0].strip() 
                        
                        if len(partes) > 1:
                            texto_fonte = f"Fonte: {partes[1].strip()}"
                        else:
                            texto_fonte = ""
                        
                        # Verificar se √© Gr√°fico ou Figura
                        eh_grafico = legenda_chave.startswith("Gr√°fico")
                        tipo_imagem = "GR√ÅFICO" if eh_grafico else "FIGURA"

                        print(f"--- Processando {tipo_imagem}: {legenda_chave}")
                        
                        # Se for gr√°fico, adicionar t√≠tulo acima
                        if eh_grafico:
                            # Extrair apenas o texto ap√≥s "Gr√°fico X -"
                            # Ex: "Gr√°fico 11 - Percentual de Magistrados..." ‚Üí "Percentual de Magistrados..."
                            match_titulo = re.match(r'^Gr√°fico\s+\d+\s*[-‚Äì]\s*(.+)', legenda_chave, re.IGNORECASE)
                            if match_titulo:
                                titulo_texto = match_titulo.group(1).strip()
                            else:
                                titulo_texto = legenda_chave  # Fallback se n√£o encontrar o padr√£o
                            
                            p_titulo = document.add_paragraph()
                            run_titulo = p_titulo.add_run(titulo_texto)
                            run_titulo.font.name = 'Calibri'
                            run_titulo.font.size = Pt(18)
                            run_titulo.bold = True
                            p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_titulo.paragraph_format.space_before = Pt(12)
                            p_titulo.paragraph_format.space_after = Pt(6)
                        
                        # Busca de imagem: 1¬∫ no mapeamento autom√°tico, 2¬∫ no MAPA_IMAGENS manual
                        caminho_imagem_abs = None
                        largura_cm = 16.5  # Largura padr√£o
                        recuo_cm = 0.0     # Recuo padr√£o (sem recuo)
                        
                        # Prioridade 1: Buscar no mapeamento autom√°tico (gerado por match_graficos.py)
                        caminho_imagem_abs = buscar_caminho_grafico(legenda_chave)
                        
                        # Prioridade 2: Buscar no MAPA_IMAGENS manual (report_data.py)
                        if not caminho_imagem_abs and legenda_chave in MAPA_IMAGENS:
                            imagem_info = MAPA_IMAGENS[legenda_chave]
                            
                            if isinstance(imagem_info, dict):
                                caminho_imagem_relativo = imagem_info.get("caminho")
                                largura_cm = imagem_info.get("width", 16.5)
                                recuo_cm = imagem_info.get("indent", 0.0)  # Novo: suporte a recuo
                            else:
                                caminho_imagem_relativo = imagem_info
                                largura_cm = 16.5
                                recuo_cm = 0.0
                            
                            caminho_imagem_abs = os.path.join(SCRIPT_DIR, caminho_imagem_relativo)
                        
                        # Inserir imagem se encontrada
                        if caminho_imagem_abs and os.path.exists(caminho_imagem_abs):
                            try:
                                document.add_picture(caminho_imagem_abs, width=Cm(largura_cm)) 
                                p_imagem = document.paragraphs[-1]
                                p_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # Aplicar recuo personalizado se definido
                                if recuo_cm != 0.0:
                                    aplicar_recuo_paragrafo(p_imagem, recuo_cm)
                                    print(f"   ‚úÖ Imagem inserida com recuo de {recuo_cm}cm: {os.path.basename(caminho_imagem_abs)}")
                                else:
                                    print(f"   ‚úÖ Imagem inserida: {os.path.basename(caminho_imagem_abs)}")
                                
                            except Exception as e:
                                print(f"   ‚ùå ERRO ao inserir imagem: {e}")
                                document.add_paragraph(f"[ERRO AO INSERIR IMAGEM: {legenda_chave}]")
                        else:
                            print(f"   ‚ö†Ô∏è  AVISO: Imagem n√£o encontrada para '{legenda_chave}'")
                            document.add_paragraph(f"[IMAGEM N√ÉO ENCONTRADA: {legenda_chave}]")

                        # Para figuras, adicionar legenda abaixo
                        # Para gr√°ficos, n√£o adicionar legenda (j√° tem t√≠tulo acima)
                        if not eh_grafico:
                            p_legenda = document.add_paragraph()
                            run_legenda = p_legenda.add_run(legenda_chave)
                            run_legenda.font.name = 'Calibri'
                            run_legenda.font.size = Pt(8) 
                            run_legenda.bold = False 
                            p_legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                            p_legenda.paragraph_format.space_before = Pt(6) 
                            p_legenda.paragraph_format.space_after = Pt(0)

                        if texto_fonte:
                            p_fonte = document.add_paragraph()
                            run_fonte = p_fonte.add_run(texto_fonte)
                            run_fonte.font.name = 'Calibri'
                            run_fonte.font.size = Pt(8) 
                            p_fonte.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            p_fonte.paragraph_format.space_after = Pt(24)

                    elif bloco['tipo'] == 'TABELA_ATOS':
                        print(f"--- Inserindo Tabela 01 (Atos) para {titulo_chave} ---")
                        adicionar_tabela_atos(document, bloco['dados'])
                    
                    elif bloco['tipo'] == 'TABELA_AREAS':
                        print(f"--- Inserindo Tabela 02 (√Åreas) para {titulo_chave} ---")
                        adicionar_tabela_areas(document, bloco['dados'])
                    
                    elif bloco['tipo'] == 'TABELA_ESTRUTURA':
                        print(f"--- Inserindo Tabela 03 (Estrutura) para {titulo_chave} ---")
                        adicionar_tabela_estrutura(document, bloco['dados'])

                    elif bloco['tipo'] == 'TABELA_COMARCAS':
                        print(f"--- Inserindo Tabela 04 (Comarcas) para {titulo_chave} ---")
                        adicionar_tabela_comarcas(document, bloco['dados'])
                    
                    elif bloco['tipo'] == 'TABELA_NUCLEOS':
                        print(f"--- Inserindo Tabela 05 (N√∫cleos) para {titulo_chave} ---")
                        adicionar_tabela_nucleos(document, bloco['dados'])

                    elif bloco['tipo'] == 'TABELA_PROCESSOS':
                        print(f"--- Inserindo Tabela 06 (Dados Hist√≥ricos) para {titulo_chave} ---")
                        adicionar_tabela_processos(document, bloco['dados'])

                    elif bloco['tipo'] == 'TABELA_JULGAMENTOS':
                        print(f"--- Inserindo Tabela 07 (Dados Hist√≥ricos) para {titulo_chave} ---")
                        adicionar_tabela_processos(document, bloco['dados'])
                        
                    elif bloco['tipo'] == 'TABELA_ACERVO':
                        print(f"--- Inserindo Tabela 08 (Dados Hist√≥ricos) para {titulo_chave} ---")
                        adicionar_tabela_processos(document, bloco['dados'])
                    
                    elif bloco['tipo'] == 'TABELA_ORCAMENTO':
                        print(f"--- Inserindo Tabela 09 (Or√ßamento) para {titulo_chave} ---") 
                        adicionar_tabela_orcamento(
                            document, 
                            TITULO_TABELA_ORCAMENTO, 
                            bloco['dados']
                        )
                    
                    elif bloco['tipo'] == 'TABELA_ORCAMENTO_ACAO':
                        print(f"--- Inserindo Tabela 10 (Or√ßamento - A√ß√£o) para {titulo_chave} ---")
                        adicionar_tabela_orcamento(
                            document, 
                            TITULO_TABELA_ORCAMENTO_ACAO,
                            bloco['dados']
                        )

                    elif bloco['tipo'] == 'TABELA_ORCAMENTO_CONJUNTO':
                        print(f"--- Inserindo Tabela 11 (Or√ßamento 2025) para {titulo_chave} ---")
                        adicionar_tabela_orcamento_conjunto( 
                            document, 
                            bloco['dados'] 
                        )
                    
                    elif bloco['tipo'] == 'TABELA_CIDADES':
                        print(f"--- Inserindo Tabela 12 (Cidades) para {titulo_chave} ---")
                        adicionar_tabela_cidades(document, bloco['dados'])

                    elif bloco['tipo'] == 'TABELA_JUSTICA_NUMEROS':
                        print(f"--- Inserindo Tabela 13 (Justi√ßa em N√∫meros) para {titulo_chave} ---")
                        adicionar_tabela_justica_numeros(document, bloco['dados'])
                        
    # --- FIM DA SE√á√ÉO ---

    # --- SALVAR O DOCUMENTO ---
    try:
        document.save(CAMINHO_SAIDA)
        print(f"Documento '{CAMINHO_SAIDA}' gerado com sucesso!")
    except PermissionError:
        print(f"!!! ERRO DE PERMISS√ÉO: N√£o foi poss√≠vel salvar o arquivo '{CAMINHO_SAIDA}'.")
        print("!!! Verifique se o arquivo n√£o est√° aberto no Word.")
    except Exception as e:
        print(f"!!! ERRO INESPERADO AO SALVAR: {e}")
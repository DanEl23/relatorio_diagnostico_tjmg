
import os
from google.colab import drive
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

style_h1 = document.styles['Heading 1']
font = style_h1.font
font_h1 = style_h1.font
font_h1.name = 'Calibri'
font_h1.size = Pt(20)
font_h1.color.rgb = RGBColor(162, 22, 18)
font_h1.all_caps = True


# Seção Principal (H1) - Usamos level=1
document.add_heading('Relatório de Análise Automática', level=1)
document.add_paragraph('Este relatório foi gerado automaticamente pelo script Python e contém análises de dados estruturados')

# Seção Primária (H2) - Usamos level=2
document.add_heading('1. Estrutura e Metodologia', level=2)

#Subseção (H3) - Usamos level=3
document.add_heading('1.1 Fonte de Dados e Parametrização', level=3)
document.add_paragraph('Os dados de entrada foram processados antes da inclusão no documento final')

document.save('Relatório_Final_Automático.docx')
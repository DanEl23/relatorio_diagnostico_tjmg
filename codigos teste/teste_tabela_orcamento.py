import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- 1. DADOS DE TESTE (SIMULAÇÃO DE report_data.py) ---
# A Tabela 09 espera 7 colunas no dicionário, usando as 2 primeiras.
dados_tabela_orcamento = [
    ("HEADER_MERGE", "Unidade Orçamentária 1031 – TJMG | Despesa Realizada por Ação Orçamentária – 2024", "", "", "", "", "", ""),
    ("SUB_HEADER", "AÇÃO ORÇAMENTÁRIA", "DESPESA REALIZADA 2024 (R$)", "", "", "", "", ""),
    ("DATA_ROW", "7004 - Precatórios e Sentenças Judiciárias", "-", "", "", "", "", ""),
    ("DATA_ROW", "7006 - Proventos de Inativos Civis e Pensionistas", "2.535.040.959,40", "", "", "", "", ""),
    ("DATA_ROW", "2053 - Remuneração de Magistrados da Ativa E Encargos Sociais", "1.353.944.848,00", "", "", "", "", ""),
    ("DATA_ROW", "2054 - Remuneração de Servidores da Ativa e Encargos Sociais", "5.448.469.921,18", "", "", "", "", ""),
    ("TOTAL_ROW", "TOTAL", "9.337.455.728,58", "", "", "", "", "")
]


# --- 2. FUNÇÕES AUXILIARES NECESSÁRIAS ---

def set_row_height_exact(row, height_twips):
    """ Define a altura exata da linha usando XML (twips). """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact') 
    
    for existing_trHeight in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing_trHeight)
        
    trPr.append(trHeight)

def set_cell_bottom_border(cell):
    """ Adiciona uma borda inferior sólida (preta, 0.5pt) a uma célula específica. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    bottom_border = OxmlElement('w:bottom') 
    bottom_border.set(qn('w:val'), 'single') 
    bottom_border.set(qn('w:sz'), '4')       
    bottom_border.set(qn('w:color'), '000000') 
    
    tcBorders.append(bottom_border)
    tcPr.append(tcBorders)

def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma célula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) 
    tcPr.append(vAlign)

def configurar_margens(documento, superior_cm, esquerda_cm, direita_cm, inferior_cm):
    """ Define as margens da seção principal do documento em centímetros. """
    section = documento.sections[0]
    
    section.top_margin = Cm(superior_cm)
    section.left_margin = Cm(esquerda_cm)
    section.right_margin = Cm(direita_cm)
    section.bottom_margin = Cm(inferior_cm)


# --- 3. FUNÇÃO DEDICADA (TABELA 09 - ORÇAMENTO) ---
def adicionar_tabela_orcamento(document, dados):
    """
    Cria a Tabela 09 (Orçamento - 2 colunas) com formatação específica,
    usando helpers de XML e sombreamento.
    """
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_HEADER_PRINCIPAL = '44546A' 
    COR_SUB_HEADER_COLUNA = '44546A'
    COR_LINHA_ZEBRADA = 'D9D9D9'
    COR_BRANCO_RGB = RGBColor(255, 255, 255)
    COR_PRETO_RGB = RGBColor(0, 0, 0)
    TAMANHO_FONTE_PADRAO = Pt(12)
    FONTE = 'Calibri'
    NUM_COLUNAS_DADOS = 2  
    
    ALTURA_LINHA_TWIPS = 284
    
    # --- ESTRUTURA E LARGURA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS_DADOS)
    
    try:
        table.style = 'Normal Table' 
    except KeyError:
        table.style = 'Table Grid'
    
    # Larguras das 2 colunas (60% para Ação, 40% para Despesa)
    largura_total_cm = 16.0 
    table.columns[0].width = Cm(largura_total_cm * 0.60)
    table.columns[1].width = Cm(largura_total_cm * 0.40)

    legenda_padrao = "Tabela 09 - Despesa realizada por ação (Unidade 1031). Fonte: Armazém de Informações - BO SIAFI/MG"
    
    data_row_index = 0

    # --- PREENCHIMENTO DA TABELA ---
    for i, row_data_full in enumerate(dados):
        tipo = row_data_full[0]
        dados_da_linha = row_data_full[1:1 + NUM_COLUNAS_DADOS] 
        
        row = table.add_row()
        set_row_height_exact(row, ALTURA_LINHA_TWIPS) 

        # Contagem para o zebrado
        if tipo == "DATA_ROW" or tipo == "TOTAL_ROW":
            data_row_index += 1
        
        for col_idx in range(NUM_COLUNAS_DADOS):
            cell = row.cells[col_idx]
            
            set_cell_vertical_alignment(cell, 'center')
            
            p = cell.paragraphs[0]
            # Coluna de valor (1) centralizada, Coluna de descrição (0) à esquerda
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            cell.text = "" 
            run = p.add_run(dados_da_linha[col_idx])
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            run.font.color.rgb = COR_PRETO_RGB 

            # --- Estilos de Linha/Célula ---
            
            # TIPO 1: HEADER MERGE (Mesclado)
            if tipo == "HEADER_MERGE":
                if col_idx == 0:
                    cell.merge(row.cells[1])
                    
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_HEADER_PRINCIPAL)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
                    run.font.color.rgb = COR_BRANCO_RGB
                    run.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # TIPO 2: SUB-HEADER (Cabeçalhos das 2 colunas)
            elif tipo == "SUB_HEADER":
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_SUB_HEADER_COLUNA)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                run.font.color.rgb = COR_BRANCO_RGB
                run.bold = True
                
                set_cell_bottom_border(cell) 
                
            # TIPO 3 & 4: DADOS (DATA_ROW e TOTAL_ROW)
            elif tipo == "DATA_ROW" or tipo == "TOTAL_ROW":
                
                if tipo == "TOTAL_ROW":
                    run.bold = True
                
                # Shading (Zebrado e Total)
                if (tipo == "DATA_ROW" and (data_row_index % 2) != 0) or tipo == "TOTAL_ROW":
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)


    # --- LEGENDA/FONTE (Tabela 09) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run(legenda_padrao)
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    
    # Espaçamento final após a tabela (15pt)
    p_titulo_tabela.paragraph_format.space_after = Pt(15)


# --- 4. BLOCO EXECUTOR DO TESTE ---

# ... (Funções auxiliares e adicionar_tabela_orcamento) ...

# --- 4. BLOCO EXECUTOR DO TESTE ---

if __name__ == "__main__":
    
    print("--- INICIANDO TESTE ISOLADO DA TABELA 09 (Orçamento - 2 Colunas) ---")
    
    # 1. DIAGNÓSTICO RÁPIDO: Verificar se os dados de teste foram carregados
    if not dados_tabela_orcamento or dados_tabela_orcamento[0][0] != "HEADER_MERGE":
        print("!!! ERRO CRÍTICO: A variável 'dados_tabela_orcamento' está vazia ou no formato incorreto.")
        print("!!! Certifique-se de que a variável está definida no topo do script de teste.")
        exit()
    
    documento_teste = Document()
    
    # Margens: Superior: 3cm, Esquerda: 3cm, Direita: 2cm, Inferior: 2cm
    configurar_margens(documento_teste, 3.0, 3.0, 2.0, 2.0)
    
    documento_teste.add_heading('Teste de Tabela 09 - Execução Orçamentária', level=1)
    
    # 2. Chama a função principal da tabela
    adicionar_tabela_orcamento(documento_teste, dados_tabela_orcamento)
    
    nome_arquivo = 'teste_tabela_orcamento.docx'
    
    try:
        documento_teste.save(nome_arquivo)
        print(f"Documento de teste '{nome_arquivo}' gerado com sucesso!")
        print("VERIFIQUE: Tabela de 2 colunas com layout de cabeçalho duplo, zebrado e espaçamento de 15pt final.")
    except PermissionError:
        print(f"!!! ERRO: Não foi possível salvar o arquivo. Verifique se '{nome_arquivo}' não está aberto no Word.")
    
    print("--- TESTE CONCLUÍDO ---")
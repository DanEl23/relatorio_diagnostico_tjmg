import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- 1. IMPORTAÇÃO DOS DADOS EXTERNOS ---
# (Certifique-se de que report_data.py está salvo com dados_tabela_comarcas)
try:
    from report_data import dados_tabela_comarcas
except ImportError:
    print("!!! ERRO CRÍTICO: Não foi possível encontrar 'dados_tabela_comarcas' em 'report_data.py'.")
    exit()

# --- 2. FUNÇÕES HELPER ---

def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma célula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) 
    tcPr.append(vAlign)

# --- 3. FUNÇÃO DEDICADA (TABELA 04 - COMARCAS) ---

def adicionar_tabela_comarcas(document, dados):
    """
    Cria a Tabela 04 (Comarcas) com 4 colunas e cabeçalho mesclado.
    (FINALIZADO com correção de cor para linhas zebradas).
    """
    
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_HEADER_MAIN = '7F7F7F'    
    # --- CORREÇÃO AQUI ---
    COR_LINHA_ZEBRADA = 'D9D9D9' 
    # --- FIM DA CORREÇÃO ---
    
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) 
    FONTE = 'Calibri'
    
    table = document.add_table(rows=0, cols=4) 
    
    try:
        table.style = 'Normal Table'
    except KeyError:
        print("Aviso: Estilo 'Normal Table' não encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    col_widths = [Cm(4.375), Cm(4.375), Cm(4.375), Cm(4.375)] 
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        row = table.add_row()
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cells = row.cells 

        # --- TIPO 1: Cabeçalho Principal (Mesclado) ---
        if tipo_linha == "HEADER_MERGE_4":
            col1_texto = row_data[1]
            cell = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
            cell.text = col1_texto
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_MAIN)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO

        # --- TIPO 2: Dados (4 Colunas) ---
        elif tipo_linha == "DATA_4_COL":
            data_row_index += 1
            
            for j in range(4):
                cell = cells[j]
                cell.text = row_data[j+1] 
                
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                
                run = p.runs[0]
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = False

                # --- APLICA SOMBREADO NA CÉLULA (tcPr) ---
                if data_row_index % 2 != 0: 
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- LEGENDA/FONTE (Tabela 04) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 04 - Comarcas Instaladas. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12)


# --- 4. BLOCO EXECUTOR DO TESTE ---

if __name__ == "__main__":
    
    print("--- INICIANDO TESTE ISOLADO DA TABELA 04 (Comarcas - FINALIZADO) ---")
    
    documento_teste = Document()
    documento_teste.add_heading('Teste de Tabela 04 - Comarcas', level=1)
    
    adicionar_tabela_comarcas(documento_teste, dados_tabela_comarcas)
    
    nome_arquivo = 'teste_tabela_comarcas_finalizada.docx'
    
    try:
        documento_teste.save(nome_arquivo)
        print(f"Documento de teste '{nome_arquivo}' gerado com sucesso!")
        print("VERIFIQUE: As linhas de dados devem ter cores alternadas (zebradas) com a cor 'D9D9D9'.")
    except PermissionError:
        print(f"!!! ERRO: Não foi possível salvar o arquivo. Verifique se '{nome_arquivo}' não está aberto no Word.")
    
    print("--- TESTE CONCLUÍDO ---")
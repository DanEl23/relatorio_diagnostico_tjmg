import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- 1. IMPORTAÇÃO DOS DADOS EXTERNOS ---
# (Importa os dados que acabamos de adicionar ao report_data.py)
try:
    from report_data import dados_tabela_estrutura
except ImportError:
    print("!!! ERRO CRÍTICO: Não foi possível encontrar 'dados_tabela_estrutura' em 'report_data.py'.")
    print("!!! Certifique-se de que 'report_data.py' está salvo e no mesmo diretório.")
    exit()

# --- 2. FUNÇÕES HELPER (Copiadas do script principal) ---

def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma célula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) 
    tcPr.append(vAlign)

def set_group_top_border(cell):
    """ Adiciona uma borda superior sólida (preta, 0.5pt) a uma célula específica. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    top_border = OxmlElement('w:top') 
    top_border.set(qn('w:val'), 'single') 
    top_border.set(qn('w:sz'), '4')       
    top_border.set(qn('w:color'), '000000') 
    
    tcBorders.append(top_border)
    tcPr.append(tcBorders)

# --- 3. NOVA FUNÇÃO DEDICADA (TABELA 03 - ESTRUTURA) ---

def adicionar_tabela_estrutura(document, dados):
    """
    Cria a Tabela 03 (Estrutura) com 1 coluna, mas mantendo a estética
    da Tabela 02 (Grupos, Bordas, Cores).
    """
    
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_HEADER_MAIN = '7F7F7F'    
    COR_HEADER_GROUP = 'D9D9D9'  
    COR_LINHA_ZEBRADA = 'EEEEEE'
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) # (Usando Pt(11) como você ajustou)
    FONTE = 'Calibri'
    
    # --- LÓGICA DE 1 COLUNA ---
    table = document.add_table(rows=0, cols=1) 
    
    try:
        table.style = 'Normal Table'
    except KeyError:
        print("Aviso: Estilo 'Normal Table' não encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    # Define a largura da coluna única (14.5 + 3.0 = 17.5cm, total da Tabela 02)
    col_widths = [Cm(17.5)] 
    table.columns[0].width = col_widths[0]

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        col1_texto = row_data[1]
        
        row = table.add_row()
        
        # Propriedades de Quebra de Página
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cell = row.cells[0] # Apenas 1 célula
        cell.text = col1_texto
        
        # --- APLICA ESTILO BASEADO NO TIPO ---
        
        set_cell_vertical_alignment(cell, 'center')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Alinhamento esquerdo como padrão
        
        # Espaçamento 1.0 (Simples)
        p_format = p.paragraph_format
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        
        run = p.runs[0]
        run.font.name = FONTE
        run.font.size = TAMANHO_FONTE_PADRAO
        run.font.color.rgb = COR_PRETO_RGB
        run.bold = False

        # --- TIPO 1: Cabeçalho Principal (Cinza Escuro) ---
        if tipo_linha == "HEADER_MAIN":
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_MAIN)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            
        # --- TIPO 2: Cabeçalho de Grupo (Cinza Claro) ---
        elif tipo_linha == "HEADER_GROUP_MERGED":
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            run.bold = True
            set_group_top_border(cell) # Adiciona a borda superior

        # --- TIPO 3: Dados (Zebrado) ---
        elif tipo_linha == "DATA_MERGED":
            data_row_index += 1
            if data_row_index % 2 == 0: # Aplica linha zebrada
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- LEGENDA/FONTE (Tabela 03) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    # (Texto da Legenda Tabela 03)
    run_titulo = p_titulo_tabela.add_run("Tabela 03 - Estruturas para a Prestação Jurisdicional na Segunda Instância. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12)


# --- 4. BLOCO EXECUTOR DO TESTE ---

if __name__ == "__main__":
    
    print("--- INICIANDO TESTE ISOLADO DA TABELA 03 (Estrutura - 1 Coluna) ---")
    
    documento_teste = Document()
    documento_teste.add_heading('Teste de Tabela 03 - Estrutura', level=1)
    
    # Chama a nova função com os dados importados
    adicionar_tabela_estrutura(documento_teste, dados_tabela_estrutura)
    
    nome_arquivo = 'teste_tabela_estrutura.docx'
    
    try:
        documento_teste.save(nome_arquivo)
        print(f"Documento de teste '{nome_arquivo}' gerado com sucesso!")
        print("Verifique o arquivo: Tabela de 1 coluna com estilo de Grupo (cor e borda superior).")
    except PermissionError:
        print(f"!!! ERRO: Não foi possível salvar o arquivo. Verifique se '{nome_arquivo}' não está aberto no Word.")
    
    print("--- TESTE CONCLUÍDO ---")
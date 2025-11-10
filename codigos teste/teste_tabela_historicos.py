import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm # Cm é essencial para margens
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- 1. IMPORTAÇÃO DOS DADOS EXTERNOS ---
try:
    from report_data import dados_tabela_historicos
except ImportError:
    print("!!! ERRO CRÍTICO: Não foi possível encontrar 'dados_tabela_historicos' em 'report_data.py'.")
    exit()

# --- 2. FUNÇÕES HELPER (Para Borda, Alinhamento e Margens) ---

def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma célula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) 
    tcPr.append(vAlign)

def set_cell_bottom_border(cell):
    """ Adiciona uma borda inferior sólida (preta, 0.5pt) a uma célula específica. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    bottom_border = OxmlElement('w:bottom') 
    bottom_border.set(qn('w:val'), 'single') 
    bottom_border.set(qn('w:sz'), '4')       
    bottom_border.set(qn('w:color'), '000000') 
    
    tcBorders.append(bottom_border)
    tcPr.append(tcBorders)
    
def configurar_margens(documento, superior_cm, esquerda_cm, direita_cm, inferior_cm):
    """ Define as margens da seção principal do documento em centímetros. """
    section = documento.sections[0]
    
    section.top_margin = Cm(superior_cm)
    section.left_margin = Cm(esquerda_cm)
    section.right_margin = Cm(direita_cm)
    section.bottom_margin = Cm(inferior_cm)

def set_row_height_exact(row, height_twips):
    """ Define a altura exata da linha usando XML (twips). """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact') 
    
    for existing_trHeight in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing_trHeight)
        
    trPr.append(trHeight)

# --- 3. FUNÇÃO DEDICADA (TABELA 06 - HISTÓRICOS) ---

def adicionar_tabela_historicos(document, dados):
    """
    Cria a Tabela 06 (Dados Históricos), aplicando cor de destaque à
    coluna do ano mais recente (2024).
    """
    
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_HEADER_PRINCIPAL = '7F7F7F'    
    COR_LINHA_ZEBRADA = 'D9D9D9'       
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(12) 
    FONTE = 'Calibri'
    NUM_COLUNAS = 7
    
    # NOVAS CORES (Hexadecimais)
    COR_SUB_HEADER_COLUNA = '44546A' # RGB(68,84,106)
    COR_DADOS_COLUNA = 'D5DCE4'      # RGB(213,220,228)
    COLUNA_DESTAQUE_INDEX = 5        # Coluna "2024" (Instância=0, 2020=1, ..., 2024=5)
    
    ALTURA_LINHA_TWIPS = 284 
    
    # --- ESTRUTURA E LARGURA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS) 
    
    try:
        table.style = 'Normal Table' 
    except KeyError:
        print("Aviso: Estilo 'Normal Table' não encontrado. Revertendo para 'Table Grid'.")
        table.style = 'Table Grid' 
    
    col_widths = [Cm(3.2)] + [Cm(2.133)] * 6
    
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    data_row_index = 0 
 
    for i, row_data_full in enumerate(dados):
        
        tipo_linha = row_data_full[0]
        row_data = row_data_full[1:] 
        
        row = table.add_row()
        set_row_height_exact(row, ALTURA_LINHA_TWIPS)
        
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha in ["HEADER_MERGE", "SUB_HEADER"]:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cells = row.cells 

        # --- TIPO 1: CABEÇALHO PRINCIPAL (Merscla 7 Colunas) ---
        if tipo_linha == "HEADER_MERGE":
            
            cell = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4]).merge(cells[5]).merge(cells[6])
            cell.text = row_data[0]
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_PRINCIPAL)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = Pt(12) 
        
        # --- TIPO 2: SUB-CABEÇALHO (Instância, 2020...) ---
        elif tipo_linha == "SUB_HEADER":
            
            for j in range(NUM_COLUNAS):
                cell = cells[j]
                cell.text = row_data[j]
                
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = True
                run.font.name = FONTE
                run.font.size = Pt(12) 
                
                # --- Destaque: Coluna 2024 ---
                if j == COLUNA_DESTAQUE_INDEX:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_SUB_HEADER_COLUNA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    run.font.color.rgb = COR_BRANCO_RGB # Texto branco para fundo escuro
                
                set_cell_bottom_border(cell)

        # --- TIPO 3 & 4: LINHAS DE DADOS E TOTAL ---
        elif tipo_linha in ["DATA_ROW", "TOTAL_ROW"]:
            
            is_total_row = (tipo_linha == "TOTAL_ROW")
            if not is_total_row:
                data_row_index += 1 

            for j, cell_data in enumerate(row_data):
                cell = cells[j]
                cell.text = cell_data
                
                cell_align = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                p.alignment = cell_align
                
                run = p.runs[0]
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = is_total_row
                
                # --- Lógica de Sombreamento (Prioridade: Coluna > Total/Zebrado) ---
                current_shading_color = None

                if is_total_row or (data_row_index % 2 != 0): 
                    # Cor de fundo padrão da linha (Total ou Zebrado Ímpar)
                    current_shading_color = COR_LINHA_ZEBRADA 
                
                if j == COLUNA_DESTAQUE_INDEX:
                    # Sobrescreve para a cor da Coluna 2024
                    current_shading_color = COR_DADOS_COLUNA
                
                if current_shading_color:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), current_shading_color)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                # --- Fim da Lógica de Sombreamento ---

                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                

    # --- LEGENDA/FONTE (Tabela 06) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 06 - Número de processos distribuídos. Fonte: Centro de Informações para a Gestão Institucional – CEINFO")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12)


# --- 4. BLOCO EXECUTOR DO TESTE ---

if __name__ == "__main__":
    
    print("--- INICIANDO TESTE ISOLADO DA TABELA 06 (Margens, Altura e Destaque de Coluna) ---")
    
    documento_teste = Document()
    
    # Margens: Superior: 3cm, Esquerda: 3cm, Direita: 2cm, Inferior: 2cm
    configurar_margens(documento_teste, 3.0, 3.0, 2.0, 2.0)
    
    documento_teste.add_heading('Teste de Tabela 06 - Processos Distribuídos', level=1)
    
    adicionar_tabela_historicos(documento_teste, dados_tabela_historicos)
    
    nome_arquivo = 'teste_tabela_historicos_final_destaque.docx'
    
    try:
        documento_teste.save(nome_arquivo)
        print(f"Documento de teste '{nome_arquivo}' gerado com sucesso!")
        print("VERIFIQUE: A coluna '2024' deve ter a cor #44546A no subcabeçalho e #D5DCE4 nas linhas de dados.")
    except PermissionError:
        print(f"!!! ERRO: Não foi possível salvar o arquivo. Verifique se '{nome_arquivo}' não está aberto no Word.")
    
    print("--- TESTE CONCLUÍDO ---")
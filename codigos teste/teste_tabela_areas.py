import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
# --- 1. GARANTA QUE WD_LINE_SPACING ESTÁ IMPORTADO ---
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- FUNÇÕES HELPER (Inalteradas, exceto pela remoção de 'limpar_espacamento_lista') ---

def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma célula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) 
    tcPr.append(vAlign)

def set_group_top_border(cell):
    """ Adiciona uma borda superior sólida (preta, 0.5pt) a uma célula específica. """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    top_border = OxmlElement('w:top') 
    top_border.set(qn('w:val'), 'single') 
    top_border.set(qn('w:sz'), '4')       
    top_border.set(qn('w:color'), '000000') 
    
    tcBorders.append(top_border)
    tcPr.append(tcBorders)

# --- 1. DADOS BRUTOS (Inalterado) ---
dados_tabela_areas = [
    # Tipo, Col 1, Col 2
    ("HEADER_MAIN", "DENOMINAÇÃO", ""),
    ("DATA_MERGED", "Comitê Estratégico de Gestão Institucional", ""),
    ("DATA_MERGED", "Comitê Gestor de Segurança da Informação", ""),
    ("DATA_MERGED", "Comitê Institucional de Inteligência", ""),
    ("DATA_MERGED", "Comitê de Governança e Gestão EstratégICA", ""), # Corrigido (se necessário)
    ("DATA_MERGED", "Comitê de Monitoramento e Suporte à Prestação Jurisdicional", ""),
    ("DATA_MERGED", "Comitê de Tecnologia da Informação e Comunicação", ""),
    
    ("HEADER_GROUP_SIGLA", "SUPERINTENDÊNCIA ADMINISTRATIVA", "SIGLA"),
    ("DATA_SPLIT", "Diretoria Executiva de Administração de Recursos Humanos", "DEARHU"),
    ("DATA_SPLIT", "Diretoria Executiva de Comunicação", "DIRCOM"),
    ("DATA_SPLIT", "Diretoria Executiva de Engenharia e Gestão Predial", "DENGEP"),
    ("DATA_SPLIT", "Diretoria Executiva de Finanças e Execução Orçamentária", "DIRFIN"),
    ("DATA_SPLIT", "Diretoria Executiva de Gestão de Bens, Serviços e Patrimônio", "DIRSEP"),
    ("DATA_SPLIT", "Diretoria Executiva de Informática", "DIRTEC"),
    ("DATA_SPLIT", "Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional", "DEPLAG"),
    ("DATA_SPLIT", "Gabinete de Segurança Institucional", "GSI"),
    ("DATA_SPLIT", "Secretaria de Auditoria Interna", "SECAUD"),
    ("DATA_SPLIT", "Secretaria de Governança e Gestão Estratégica", "SEGOVE"),
    ("DATA_SPLIT", "Secretaria do Órgão Especial", "SEOESP"),

    ("HEADER_GROUP_MERGED", "SUPERINTENDÊNCIA DO 1º VICE-PRESIDENTE", ""),
    ("DATA_SPLIT", "Diretoria Executiva de Suporte à Prestação Jurisdicional", "DIRSUP"),
    ("DATA_SPLIT", "Secretaria de Padronização e Acompanhamento da Gestão Judiciária", "SEPAD"),

    ("HEADER_GROUP_MERGED", "SUPERINTENDÊNCIA DO 2º VICE-PRESIDENTE", ""),
    ("DATA_SPLIT", "Diretoria Executiva de Desenvolvimento de Pessoas", "DIRDEP"),
    ("DATA_SPLIT", "Diretoria Executiva de Gestão da Informação Documental", "DIRGED"),

    ("HEADER_GROUP_MERGED", "SUPERINTENDÊNCIA DO 3º VICE-PRESIDENTE", ""),
    ("DATA_SPLIT", "Assessoria de Gestão da Inovação", "AGIN"),
    ("DATA_SPLIT", "Núcleo Permanente de Métodos Consensuais de Solução de Conflitos", "NUPEMEC"),

    ("HEADER_GROUP_MERGED", "CORREGEDORIA-GERAL DE JUSTIÇA", ""),
    ("DATA_SPLIT", "Diretoria Executiva de Atividade Correcional", "DIRCOR"),
    ("DATA_SPLIT", "Secretaria de Suporte ao Planejamento e à Gestão da Primeira Instância", "SEPLAN")
]

# --- 2. FUNÇÃO DEDICADA DA TABELA 02 (Espaçamento 1.0) ---

def adicionar_tabela_areas(document, dados):
    """
    Cria a Tabela 02 com alinhamento esquerdo, largura de 3cm, bordas de grupo
    e espaçamento de linha 1.0.
    """
    
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_HEADER_MAIN = '7F7F7F'    
    COR_HEADER_GROUP = 'D9D9D9'  
    COR_LINHA_ZEBRADA = 'EEEEEE'
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    TAMANHO_FONTE_PADRAO = Pt(11) 
    FONTE = 'Calibri'
    
    table = document.add_table(rows=0, cols=2)
    table.style = 'Normal Table' # (Usando o estilo que funcionou no seu último teste)
    
    col_widths = [Cm(14.5), Cm(3.0)] 
    table.columns[0].width = col_widths[0]
    table.columns[1].width = col_widths[1]

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        
        tipo_linha = row_data[0]
        col1_texto = row_data[1]
        col2_texto = row_data[2]
        
        row = table.add_row()
        
        # Propriedades de Quebra de Página (Mantido)
        tr = row._tr 
        trPr = tr.get_or_add_trPr() 
        if tipo_linha.startswith("HEADER"):
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') 
            trPr.append(cantSplit)
        
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        
        # --- TIPO 1: Cabeçalho Principal (DENOMINAÇÃO) ---
        if tipo_linha == "HEADER_MAIN":
            cell1.merge(cell2)
            cell = cell1
            cell.text = col1_texto
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_MAIN)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # --- Ajuste de Espaçamento 1.0 ---
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            
            run = p.runs[0]
            run.font.color.rgb = COR_BRANCO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            
        # --- TIPO 2: Cabeçalho de Grupo (com SIGLA) ---
        elif tipo_linha == "HEADER_GROUP_SIGLA":
            for j, cell in enumerate([cell1, cell2]):
                cell.text = col1_texto if j == 0 else col2_texto
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                set_cell_vertical_alignment(cell, 'center')
                p = cell.paragraphs[0]
                run = p.runs[0]
                run.font.color.rgb = COR_PRETO_RGB
                run.bold = True
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_group_top_border(cell) # Bordas Superiores

                # --- Ajuste de Espaçamento 1.0 ---
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)

        # --- TIPO 3: Cabeçalho de Grupo (Mesclado) ---
        elif tipo_linha == "HEADER_GROUP_MERGED":
            cell1.merge(cell2)
            cell = cell1
            cell.text = col1_texto
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), COR_HEADER_GROUP)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.color.rgb = COR_PRETO_RGB
            run.bold = True
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO
            
            set_group_top_border(cell) # Bordas Superiores

            # --- Ajuste de Espaçamento 1.0 ---
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)

        # --- TIPO 4: Dados (Mesclados - Comitês) ---
        elif tipo_linha == "DATA_MERGED":
            data_row_index += 1
            cell1.merge(cell2)
            cell = cell1
            cell.text = col1_texto
            
            set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.color.rgb = COR_PRETO_RGB
            run.font.name = FONTE
            run.font.size = TAMANHO_FONTE_PADRAO

            # --- Ajuste de Espaçamento 1.0 ---
            p_format = p.paragraph_format
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)

            if data_row_index % 2 == 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                cell._tc.get_or_add_tcPr().append(shading_elm)

        # --- TIPO 5: Dados (Divididos - Áreas/Siglas) ---
        elif tipo_linha == "DATA_SPLIT":
            data_row_index += 1
            
            # Formatação Célula 1
            cell1.text = col1_texto
            set_cell_vertical_alignment(cell1, 'center')
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run1 = p1.runs[0]
            run1.font.color.rgb = COR_PRETO_RGB
            run1.font.name = FONTE
            run1.font.size = TAMANHO_FONTE_PADRAO
            
            # --- Ajuste de Espaçamento 1.0 (Célula 1) ---
            p1_format = p1.paragraph_format
            p1_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p1_format.space_before = Pt(0)
            p1_format.space_after = Pt(0)

            # Formatação Célula 2
            cell2.text = col2_texto
            set_cell_vertical_alignment(cell2, 'center')
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.runs[0]
            run2.font.color.rgb = COR_PRETO_RGB
            run2.font.name = FONTE
            run2.font.size = TAMANHO_FONTE_PADRAO
            
            # --- Ajuste de Espaçamento 1.0 (Célula 2) ---
            p2_format = p2.paragraph_format
            p2_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p2_format.space_before = Pt(0)
            p2_format.space_after = Pt(0)
            
            # Linhas Zebradas
            if data_row_index % 2 == 0:
                for cell in [cell1, cell2]:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), COR_LINHA_ZEBRADA)
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- LEGENDA/FONTE (Inalterado) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 02 - Principais áreas da Secretaria do TJMG. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12) 


# --- 3. BLOCO EXECUTOR DO TESTE ---

if __name__ == "__main__":
    
    print("--- INICIANDO TESTE ISOLADO DA TABELA 02 (Espaçamento 1.0) ---")
    
    documento_teste = Document()
    documento_teste.add_heading('Teste de Tabela 02 - Áreas da Secretaria', level=1)
    
    adicionar_tabela_areas(documento_teste, dados_tabela_areas)
    
    nome_arquivo = 'teste_tabela_areas_final_espacamento_1.docx'
    
    try:
        documento_teste.save(nome_arquivo)
        print(f"Documento de teste '{nome_arquivo}' gerado com sucesso!")
        print("Verifique: Espaçamento de linha 1.0 (Simples) em todas as células.")
    except PermissionError:
        print(f"!!! ERRO: Não foi possível salvar o arquivo. Verifique se '{nome_arquivo}' não está aberto no Word.")
    
    print("--- TESTE CONCLUÍDO ---")
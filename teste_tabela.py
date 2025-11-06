import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- FUNÇÃO HELPER PARA ALINHAMENTO VERTICAL (INALTERADO) ---
def set_cell_vertical_alignment(cell, align):
    """Define o alinhamento vertical de uma célula usando w:vAlign (XML)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align) # 'center', 'top', 'bottom'
    tcPr.append(vAlign)

# --- FUNÇÃO HELPER PARA LIMPEZA DE ESPAÇAMENTO (INALTERADO) ---
def limpar_espacamento_lista(paragraph):
    """
    Remove o espaçamento extra de um parágrafo de lista usando manipulação XML.
    """
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    
    spacing.set(qn('w:line'), '240') 
    spacing.set(qn('w:lineRule'), 'auto')
    
    if hasattr(pPr, 'spacing'):
        try:
            pPr.remove(pPr.spacing)
        except:
            pass
        
    pPr.append(spacing)

# --- 1. DADOS BRUTOS DA TABELA DE ATOS (COMPLETO) ---
dados_tabela_atos = [
    # Cabeçalho
    ("Ato Normativo", "Estrutura"),
    ("Lei Complementar nº 59/2001", "Contém a organização e a divisão judiciárias do Estado de Minas Gerais."),
    ("Resolução do Tribunal Pleno nº 03/2012", "Contém o Regimento Interno do Tribunal de Justiça."),
    ("Resolução nº 518/2007", "Dispõe sobre os níveis hierárquicos e as atribuições gerais das unidades organizacionais que integram a Secretaria do Tribunal de Justiça do Estado de Minas Gerais."),
    ("Resolução nº 522/2007", 
     "Dispõe sobre a Superintendência Administrativa:\n"
     "ü Superintendente Administrativo;\n"
     "ü Diretoria Executiva da Gestão de Bens, Serviços e Patrimônio;\n"
     "ü Diretoria Executiva de Engenharia e Gestão Predial;\n"
     "ü Diretoria Executiva de Informática."),
    ("Resolução nº 557/2008", "Dispõe sobre a criação da Comissão Estadual Judiciária de Adoção, CEJA-MG."),
    ("Resolução nº 640/2010", "Cria a Coordenadoria da Infância e da Juventude."),
    ("Resolução nº 673/2011", "Cria a Coordenadoria da Mulher em Situação de Violência Doméstica e Familiar."),
    ("Resolução nº 821/2016", "Dispõe sobre a reestruturação da Corregedoria Geral de Justiça."),
    ("Resolução nº 862/2017", "Dispõe sobre a estrutura e o funcionamento da Ouvidoria do Tribunal de Justiça do Estado de Minas Gerais."),
    ("Resolução nº 873/2018", "Dispõe sobre a estrutura e o funcionamento do Núcleo Permanente de Métodos de Solução de Conflitos, da Superintendência da Gestão de Inovação e do órgão jurisdicional da Secretaria do Tribunal de Justiça diretamente vinculado à Terceira Vice-Presidência, e estabelece normas para a instalação dos Centros Judiciários de Solução de Conflitos e Cidadania."),
    ("Resolução nº 877/2018", "Instala, \"ad referendum\" do Órgão Especial, a 19ª Câmara Cível no Tribunal de Justiça."),
    ("Resolução n° 878/2018", "Referenda a instalação da Câmara de que trata o art. 7º da Lei Complementar estadual nº 146, de 9 de janeiro de 2018, promovida pela Resolução nº 877, de 29 de junho de 2018."),
    ("Resolução nº 886/2019", "Determina a instalação da 8ª Câmara Criminal no Tribunal de Justiça."),
    ("Resolução nº 893/2019", "Determina a instalação da 20ª Câmara Cível no Tribunal de Justiça."),
    ("Resolução n° 969/2021", 
     "Dispõe sobre os Comitês de Assessoramento à Presidência, estabelece a estrutura e o funcionamento das unidades organizacionais da Secretaria do Tribunal de Justiça diretamente vinculadas ou subordinadas à Presidência:\n"
     "ü Comitê de Governança e Gestão Estratégica;\n"
     "ü Comitê Executivo de Gestão Institucional;\n"
     "ü Comitê Institucional de Inteligência;\n"
     "ü Comitê de Monitoramento e Suporte à Prestação Jurisdicional;\n"
     "ü Comitê de Tecnologia da Informação e Comunicação;\n"
     "ü Comitê Gestor de Segurança da Informação;\n"
     "ü Comitê Gestor da Política Judiciária para a Primeira Infância; (Alínea acrescentada pela Resolução do Órgão Especial nº 1052/2023).\n"
     "ü Comitê Gestor Regional de Primeira Instância. (Alínea acrescentada pela Resolução do Órgão Especial nº 1063/2023).\n"
     "ü Secretaria de Governança e Gestão Estratégica;\n"
     "ü Diretoria Executiva de Comunicação;\n"
     "ü Gabinete de Segurança Institucional;\n"
     "ü Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional;\n"
     "ü Gerência de Suporte aos Juizados Especiais;\n"
     "ü Secretaria do Órgão Especial;\n"
     "ü Assessoria de Precatórios;\n"
     "ü Secretaria de Auditoria Interna;\n"
     "ü Memória do Judiciário."),
    ("Resolução nº 971/2021", "Institui o Programa de Justiça Restaurativa e dispõe sobre a estrutura e funcionamento do Comitê de Justiça Restaurativa - COMJUR e da Central de Apoio à Justiça Restaurativa – CEAJUR."),
    ("Resolução nº 977/2021", "Determina a instalação da Vigésima Primeira Câmara Cível e da Nona Câmara Criminal, a especialização de Câmaras no Tribunal de Justiça."),
    ("Resolução nº 979/2021", "Dispõe sobre a estrutura organizacional e o regulamento da Escola Judicial Desembargador Edésio Fernandes - EJEF."),
    ("Resolução nº 1053/2023", "Dispõe sobre a Superintendência Judiciária."),
    ("Resolução nº 1062/2023", "Altera a Resolução do Órgão Especial nº 979, de 17 de novembro de 2021, que \"Dispõe sobre a estrutura organizacional e o regulamento da Escola Judicial Desembargador Edésio Fernandes - EJEF\"."),
    ("Resolução nº 1063/2023", "Dispõe sobre a organização e o funcionamento do Comitê Gestor Regional de Primeira Instância no âmbito do Poder Judiciário do Estado de Minas Gerais."),
    ("Resolução nº 1066/2023", "Dispõe sobre a estrutura e o funcionamento do Grupo de Monitoramento e Fiscalização do Sistema Carcerário e Socioeducativo - GMF no âmbito do Tribunal de Justiça do Estado de Minas Gerais."),
    ("Resolução nº 1079/2024", "Altera a Resolução do Órgão Especial nº 979, de 17 de novembro de 2021, que \"Dispõe sobre a estrutura organizacional e o regulamento da Escola Judicial Desembargador Edésio Fernandes - EJEF."),
    ("Resolução nº 1080/2024", "Institui o Regulamento da Escola Judicial Desembargador Edésio Fernandes - EJEF."),
    ("Resolução nº 1086/2024", "Altera a Resolução do Órgão Especial nº 1.010, de 29 de agosto de 2020, que \"Dispõe sobre a implementação, a estrutura e o funcionamento dos \"Núcleos de Justiça 4.0\" e dá outras providências\", e altera a Resolução do Órgão Especial nº 1.053, de 20 de setembro de 2023, que \"Dispõe sobre a Superintendência Judiciária e dá outras providências\".")
]

# --- 2. FUNÇÃO DEDICADA DE CRIAÇÃO E ESTILIZAÇÃO DA TABELA (COM QUEBRA DE PÁGINA) ---

def adicionar_tabela_atos(document, dados):
    """
    Cria e estiliza a Tabela de Atos Normativos com formatação específica,
    incluindo repetição de cabeçalho e permissão de quebra de linha.
    """
    
    # --- VARIÁVEIS DE COR E ESTILO ---
    COR_CABECALHO_RGB = RGBColor(127, 127, 127)   
    COR_CABECALHO_HEX = '7F7F7F'                  
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    COR_CINZA_CLARO_HEX = 'EEEEEE'                 

    TAMANHO_FONTE_PADRAO = Pt(12) 
    FONTE = 'Calibri'
    
    # Cria a tabela
    table = document.add_table(rows=1, cols=len(dados[0]))
    table.style = 'Table Grid'
    
    # Definir Larguras de Coluna (Cm) (Valores Corrigidos)
    col_widths = [Cm(5), Cm(16)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Processar Linhas e Estilização
    for i, row_data in enumerate(dados):
        if i > 0:
            row = table.add_row()
        else:
            row = table.rows[0]
            
        # --- INÍCIO DA NOVA LÓGICA: PROPRIEDADES DE QUEBRA DE PÁGINA ---
        tr = row._tr # Acessa o elemento XML da linha (<tr>)
        trPr = tr.get_or_add_trPr() # Acessa as propriedades da linha (<trPr>)

        if i == 0:
            # 1. REPETIR CABEÇALHO
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        else:
            # 2. PERMITIR QUEBRA DE LINHA (Desliga o 'cantSplit')
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0') # 0 = false (Permitir quebra)
            trPr.append(cantSplit)
        # --- FIM DA NOVA LÓGICA ---
            
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            
            set_cell_vertical_alignment(cell, 'center') 
            cell.text = "" 
            
            lines = cell_data.split('\n')
            is_first_content_line = True

            for k, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue 

                is_list_item = line.startswith('ü')
                
                if is_first_content_line:
                    current_paragraph = cell.paragraphs[0]
                    is_first_content_line = False
                else:
                    current_paragraph = cell.add_paragraph()

                text_to_insert = line.replace('ü', '').strip() if is_list_item else line
                run = current_paragraph.add_run(text_to_insert)

                if is_list_item:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_paragraph.style = 'List Bullet' 
                    limpar_espacamento_lista(current_paragraph)
                    
                elif i == 0:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_paragraph.paragraph_format.space_after = Pt(8) 
                    current_paragraph.paragraph_format.line_spacing = 1.0
                
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                
                if i == 0:
                    run.font.color.rgb = COR_BRANCO_RGB 
                    run.bold = True
                else:
                    run.font.color.rgb = COR_PRETO_RGB 
                    run.bold = False
            
            # 3. SOMBREADO
            if i == 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), COR_CABECALHO_HEX)
                cell._tc.get_or_add_tcPr().append(shading_elm)
            else:
                if i % 2 == 0:
                   shading_elm = OxmlElement('w:shd')
                   shading_elm.set(qn('w:fill'), COR_CINZA_CLARO_HEX)
                   cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- LEGENDA/FONTE (CORRIGIDA) ---
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 01 - Atos Normativos referentes à Estrutura do TJMG. Fonte: Portal TJMG")
    
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)
    p_titulo_tabela.paragraph_format.space_after = Pt(12) 


# --- BLOCO EXECUTOR DO TESTE ---

if __name__ == "__main__":
    
    print("--- INICIANDO TESTE ISOLADO DA TABELA (Com Repetição de Cabeçalho) ---")
    
    documento_teste = Document()
    documento_teste.add_heading('Teste de Tabela Estilizada (3.1 Atos Normativos)', level=1)
    
    # Adiciona a Tabela
    adicionar_tabela_atos(documento_teste, dados_tabela_atos)
    
    nome_arquivo = 'teste_tabela_atos_quebra_pagina.docx'
    
    try:
        documento_teste.save(nome_arquivo)
        print(f"Documento de teste '{nome_arquivo}' gerado com sucesso!")
        print("Verifique o arquivo: O cabeçalho da tabela deve se repetir nas páginas 2 e 3.")
    except PermissionError:
        print(f"!!! ERRO: Não foi possível salvar o arquivo. Verifique se '{nome_arquivo}' não está aberto no Word.")
    
    print("--- TESTE CONCLUÍDO ---")